#!/usr/bin/env python
"""
将 Markdown 转为 Word（.docx），按标题层级映射为 Word 内置「标题 1–9」样式，
便于交底书交付代理人或所内流程。

支持：ATX 标题 (#–######)、段落、**粗体**、行内 `代码`、无序/有序列表、
围栏代码块、简单 GFM 表格、引用块（>）、水平线（---）、行内图片 ``![](path.png)``
（在最大宽、最大高约束下**等比缩放**，竖图自动缩小宽度以整图落入版面）。

**连续多行正文**（中间无空行、且非列表/标题等）时，**每一行**输出为 Word 中**独立一段**，
以便「（1）…（2）…」等分条换行；若须在同一段内接排，请写**同一行**内或用 Markdown 空行分隔逻辑段。
被标题、段落、表格等隔开的 Markdown 有序列表，在 Word 中各自从 1 重计，避免跨章串号。

定稿宜先用同目录 **`mermaid_render.py`** 将 **mermaid** 转为 PNG；**LaTeX 公式**优先经 **`math_to_omml.py`**（``latex2mathml``）写入 **可编辑 Office Math**，失败则留原文。公式 PNG 须 ``--math-render``（可选 ``matplotlib``），且仅在用户确认后安装。

用法：
  python md_to_docx.py --input disclosure.md --output disclosure.docx
  python md_to_docx.py -i a.md -o b.docx --base-dir .   # 解析图片相对路径
  python md_to_docx.py -i a.md -o a.docx --no-omml      # 仅 PNG/原文（旧行为）
  python md_to_docx.py -i a.md -o a.docx --math-render  # OMML 失败时用公式 PNG（须 matplotlib）

依赖：python-docx + latex2mathml（根目录 requirements.txt）；matplotlib 为可选
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from stdio_utf8 import ensure_utf8_stdio
except ImportError:
    from tools.shared.stdio_utf8 import ensure_utf8_stdio

# 插图最大尺寸（英寸）：在常见 A4、默认边距下保证整图可见、按比例缩放（不过宽也不过高）。
_DEFAULT_IMAGE_MAX_W_IN = 5.5
_DEFAULT_IMAGE_MAX_H_IN = 8.2
# 公式图 Word 嵌入上限（英寸）：在「按 PNG 像素/渲染 DPI 的自然尺寸」基础上封顶，禁止强行拉高导致单行式巨字
_FORMULA_INLINE_MAX_H_IN = 0.22
_FORMULA_BLOCK_MAX_H_IN = 0.36
_FORMULA_BLOCK_MAX_W_IN = 5.5
# 与 math_render 默认 --dpi 对齐，用于把像素换算为英寸
_FORMULA_RENDER_DPI = 220.0
# 兼容旧名
_FORMULA_DISPLAY_MAX_H_IN = _FORMULA_INLINE_MAX_H_IN

_MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_HIDDEN_MD_IMAGE_COMMENT_RE = re.compile(
    r"<!--\s*!\[([^\]]*)\]\(([^)]+)\)\s*-->"
)
_INLINE_MATH_WITH_HIDDEN_IMG_RE = re.compile(
    r"(?<!\$)\$(?!\$)((?:\\.|[^$\n])+?)\$(?!\$)\s*"
    r"<!--\s*!\[([^\]]*)\]\(([^)]+)\)\s*-->"
)

# 模块级默认：可由 convert / CLI 覆盖
_PREFER_OMML = True
_SNIPPET_MAX = 120


@dataclass
class MathOutcomeStats:
    """一次 convert 的公式去向（仅统计尝试了 OMML 的条目）。"""

    omml: int = 0
    png: int = 0
    text: int = 0
    text_latex: list[str] = field(default_factory=list)

    def reset(self) -> None:
        self.omml = 0
        self.png = 0
        self.text = 0
        self.text_latex.clear()

    def record(self, latex: str, strategy: str) -> None:
        if not _PREFER_OMML:
            return
        if strategy == "omml":
            self.omml += 1
            return
        if strategy == "png":
            self.png += 1
            return
        if strategy == "text":
            self.text += 1
            snippet = (latex or "").strip().replace("\n", " ")
            if len(snippet) > _SNIPPET_MAX:
                snippet = snippet[: _SNIPPET_MAX - 3] + "..."
            if snippet:
                self.text_latex.append(snippet)

    def report(self) -> None:
        if self.omml + self.png + self.text == 0:
            return
        print(
            f"MATH: omml={self.omml} png={self.png} text={self.text}",
            file=sys.stderr,
        )
        for snippet in self.text_latex:
            print(f"OMML_FAIL: {snippet}", file=sys.stderr)
        if self.text:
            print(
                f"omml_text_fallback={self.text}",
                file=sys.stderr,
            )
        print(
            f"[md_to_docx] 公式：OMML {self.omml} 成功，PNG {self.png}，原文 {self.text}",
            file=sys.stderr,
        )


_MATH_STATS = MathOutcomeStats()


def get_math_stats() -> MathOutcomeStats:
    return _MATH_STATS


def _try_append_omml(paragraph, latex: str, *, display: bool) -> bool:
    """尝试把 LaTeX 挂为 OMML；成功 True。"""
    if not _PREFER_OMML or not (latex or "").strip():
        return False
    try:
        from math_to_omml import try_latex_to_omml
    except ImportError:
        try:
            from tools.shared.math_to_omml import try_latex_to_omml
        except ImportError:
            return False
    omml = try_latex_to_omml(latex, display=display)
    if omml is None:
        return False
    try:
        paragraph._p.append(omml)
        return True
    except Exception:
        return False


def _add_block_equation(
    doc: Document,
    latex: str,
    *,
    base_dir: Path | None = None,
    hidden: tuple[str, str] | None = None,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
    raw_fallback_lines: list[str] | None = None,
) -> str:
    """块级公式：OMML → PNG → 原文。返回所用策略名。"""
    latex = (latex or "").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    if latex and _try_append_omml(p, latex, display=True):
        _MATH_STATS.record(latex, "omml")
        return "omml"
    # OMML 失败：去掉空段，改 PNG / 原文
    p_element = p._element
    parent = p_element.getparent()
    if parent is not None:
        parent.remove(p_element)

    if hidden and _formula_image_kind(*hidden):
        ipath = _resolve_image_path(hidden[1], base_dir) if base_dir else None
        if ipath:
            _embed_from_image_ref(
                hidden[0],
                hidden[1],
                base_dir,
                doc=doc,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
            )
            _MATH_STATS.record(latex, "png")
            return "png"
    _add_math_fallback_block(doc, raw_fallback_lines or ([latex] if latex else [""]))
    _MATH_STATS.record(latex, "text")
    return "text"


def _parse_hidden_image_comment(line: str) -> tuple[str, str] | None:
    m = _HIDDEN_MD_IMAGE_COMMENT_RE.match(line.strip())
    if not m:
        return None
    return m.group(1), m.group(2).strip()


def _try_embed_hidden_comment_line(
    doc: Document,
    line: str,
    base_dir: Path | None,
    *,
    image_max_w_in: float,
    image_max_h_in: float,
) -> bool:
    hidden = _parse_hidden_image_comment(line)
    if not hidden or not base_dir:
        return False
    alt, src = hidden
    if not _resolve_image_path(src, base_dir):
        return False
    _embed_from_image_ref(
        alt,
        src,
        base_dir,
        doc=doc,
        image_max_w_in=image_max_w_in,
        image_max_h_in=image_max_h_in,
    )
    return True


def _image_pixel_size(path: Path) -> tuple[int, int] | None:
    """读取常见位图宽高（像素），失败返回 None。不依赖 Pillow。"""
    try:
        raw = path.read_bytes()
    except OSError:
        return None
    if len(raw) >= 24 and raw.startswith(b"\x89PNG\r\n\x1a\n") and raw[12:16] == b"IHDR":
        w = int.from_bytes(raw[16:20], "big")
        h = int.from_bytes(raw[20:24], "big")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 10 and raw[:3] == b"GIF" and raw[3:6] in (b"87a", b"89a"):
        w = int.from_bytes(raw[6:8], "little")
        h = int.from_bytes(raw[8:10], "little")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 4 and raw.startswith(b"\xff\xd8"):
        i = 2
        n = len(raw)
        while i < n:
            if raw[i] != 0xFF:
                i += 1
                continue
            i += 1
            while i < n and raw[i] == 0xFF:
                i += 1
            if i >= n:
                break
            marker = raw[i]
            i += 1
            if marker in (0xD8, 0xD9):
                continue
            if marker == 0xDA:
                break
            if 0xD0 <= marker <= 0xD7:
                continue
            if i + 2 > n:
                break
            seg_len = int.from_bytes(raw[i : i + 2], "big")
            if seg_len < 2:
                break
            i += 2
            if marker in (0xC0, 0xC1, 0xC2) and i + 5 <= n:
                h = int.from_bytes(raw[i + 1 : i + 3], "big")
                w = int.from_bytes(raw[i + 3 : i + 5], "big")
                if w > 0 and h > 0:
                    return w, h
            i += seg_len - 2
    return None


def _fit_image_display_inches(
    px_w: int,
    px_h: int,
    *,
    max_w_in: float,
    max_h_in: float,
) -> tuple[Inches, Inches]:
    """在不超过 max_w / max_h 的前提下等比缩放，使整图落入版面。"""
    if px_w <= 0 or px_h <= 0:
        return Inches(max_w_in), Inches(max_h_in * 0.5)
    aw = max_w_in
    ah = aw * px_h / px_w
    if ah > max_h_in:
        ah = max_h_in
        aw = ah * px_w / px_h
    return Inches(aw), Inches(ah)


def _formula_image_kind(alt: str, src: str) -> str | None:
    """返回 ``block`` / ``inline`` 表示公式图，否则 None（含注释内引用）。"""
    a = alt or ""
    s = src.replace("\\", "/")
    if "math_figures" not in s and "公式" not in a:
        return None
    if "行内" in a:
        return "inline"
    return "block"


def _is_diagram_image(alt: str, src: str) -> bool:
    """mermaid 系统框图 / 流程图等（非公式，用全幅插图尺寸）。"""
    a = alt or ""
    s = src.replace("\\", "/")
    if "mermaid_figures" in s:
        return True
    if a.startswith("图示") or a.startswith("图 "):
        return True
    return False


def _span_overlaps(spans: list[tuple[int, int]], start: int, end: int) -> bool:
    return any(not (end <= s or start >= e) for s, e in spans)


def _embed_from_image_ref(
    alt: str,
    src: str,
    base_dir: Path | None,
    *,
    doc: Document | None = None,
    paragraph=None,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> None:
    """按公式 / 框图 / 普通图规则嵌入 PNG（仅公式用小尺寸）。"""
    ipath = _resolve_image_path(src, base_dir) if base_dir else None
    missing = f"[图片缺失: {alt or src}]"
    if not ipath:
        if paragraph is not None:
            paragraph.add_run(missing)
        elif doc is not None:
            doc.add_paragraph().add_run(missing)
        return

    kind = _formula_image_kind(alt, src)
    if kind == "inline":
        p = paragraph
        if p is None and doc is not None:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
        if p is not None:
            _embed_picture_inline(p, ipath, max_h_in=_FORMULA_INLINE_MAX_H_IN)
        return

    if doc is None:
        if paragraph is not None:
            paragraph.add_run(missing)
        return

    if kind == "block":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(3)
        _embed_picture_inline(
            p,
            ipath,
            max_h_in=_FORMULA_BLOCK_MAX_H_IN,
            max_w_in=_FORMULA_BLOCK_MAX_W_IN,
        )
    else:
        _embed_picture(
            doc,
            ipath,
            alt=alt,
            src=src,
            max_w_in=image_max_w_in,
            max_h_in=image_max_h_in,
            center=False,
        )


def _maybe_render_math_md(md_text: str, base_dir: Path) -> str:
    """若含 LaTeX 公式则尝试调用 ``math_render``（已注释的 PNG 引用会跳过）。"""
    if not re.search(r"\$\$|\\\[|\\\(|(?<!\$)\$(?!\$)", md_text):
        return md_text
    try:
        from math_render import render_markdown_math
    except ImportError:
        print(
            "[md_to_docx] 未安装 matplotlib，公式将按原文写入 Word",
            file=sys.stderr,
        )
        return md_text
    stub = base_dir / "_md_to_docx_math_stub.md"
    new_md, ok, failed = render_markdown_math(
        md_text,
        out_md_path=stub,
        assets_rel="math_figures",
    )
    if ok or failed:
        print(
            f"[md_to_docx] 公式渲染：{ok} 成功，{failed} 保留原文",
            file=sys.stderr,
        )
    return new_md


def _add_math_fallback_block(doc: Document, lines: list[str]) -> None:
    """未渲染成功的 ``$$ ... $$`` 以等宽原文写入 Word。"""
    body = [ln.rstrip("\n") for ln in lines]
    _add_code_block(doc, ["$$", *body, "$$"])


def _embed_picture(
    doc: Document,
    path: Path,
    *,
    alt: str,
    src: str,
    max_w_in: float,
    max_h_in: float,
    center: bool,
) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    try:
        dims = _image_pixel_size(path)
        if dims:
            w_in, h_in = _fit_image_display_inches(
                *dims, max_w_in=max_w_in, max_h_in=max_h_in
            )
            run = p.add_run()
            run.font.bold = False
            run.add_picture(str(path.resolve()), width=w_in, height=h_in)
        else:
            run = p.add_run()
            run.font.bold = False
            run.add_picture(str(path.resolve()), width=Inches(max_w_in))
    except Exception:
        p.add_run(f"[图片无法嵌入: {alt or src} — {path}]")


def _embed_picture_inline(
    paragraph,
    path: Path,
    *,
    max_h_in: float,
    max_w_in: float | None = None,
    render_dpi: float = _FORMULA_RENDER_DPI,
) -> None:
    """按渲染 DPI 换算自然尺寸，再限制在 max_h / max_w 内；**不上拉**矮图。"""
    try:
        dims = _image_pixel_size(path)
        run = paragraph.add_run()
        run.font.bold = False
        if dims:
            px_w, px_h = dims
            dpi = render_dpi if render_dpi > 0 else 220.0
            nat_w = px_w / dpi
            nat_h = px_h / dpi
            w_in, h_in = _fit_image_display_inches(
                px_w,
                px_h,
                max_w_in=min(nat_w, max_w_in) if max_w_in is not None else nat_w,
                max_h_in=min(nat_h, max_h_in),
            )
            # 极矮图（旧小字号）给一个下限，避免再度变成细条；但不高于 max_h
            min_h = min(0.16, max_h_in)
            if h_in.inches < min_h and px_h > 0:
                h_in = Inches(min_h)
                w_in = Inches(min_h * px_w / px_h)
                if max_w_in is not None and w_in.inches > max_w_in:
                    w_in = Inches(max_w_in)
                    h_in = Inches(max_w_in * px_h / px_w)
            run.add_picture(str(path.resolve()), width=w_in, height=h_in)
        else:
            run.add_picture(str(path.resolve()), height=Inches(min(0.2, max_h_in)))
    except Exception:
        paragraph.add_run(f"[行内公式图缺失: {path}]")


def _iter_inline_paren_math_with_img(
    text: str,
) -> list[tuple[int, int, str, str, str]]:
    """匹配 ``\\(...\\)`` + 紧随的公式 PNG 注释；返回 (start,end,latex,alt,src)。"""
    out: list[tuple[int, int, str, str, str]] = []
    i = 0
    n = len(text)
    while i < n:
        if text.startswith("\\(", i):
            j = i + 2
            while j < n:
                if text.startswith("\\)", j):
                    latex = text[i + 2 : j]
                    after = j + 2
                    m = re.match(
                        r"\s*<!--\s*!\[([^\]]*)\]\(([^)]+)\)\s*-->",
                        text[after:],
                    )
                    if m:
                        end = after + m.end()
                        out.append(
                            (i, end, latex, m.group(1), m.group(2).strip())
                        )
                        i = end
                    else:
                        i = after
                    break
                if text[j] == "\\" and j + 1 < n:
                    j += 2
                else:
                    j += 1
            else:
                i += 2
        else:
            i += 1
    return out


def _iter_inline_paren_math_bare(text: str) -> list[tuple[int, int, str]]:
    """无 PNG 注释的 ``\\(...\\)`` → (start, end, latex)。"""
    out: list[tuple[int, int, str]] = []
    i = 0
    n = len(text)
    while i < n:
        if text.startswith("\\(", i):
            j = i + 2
            while j < n:
                if text.startswith("\\)", j):
                    out.append((i, j + 2, text[i + 2 : j]))
                    i = j + 2
                    break
                if text[j] == "\\" and j + 1 < n:
                    j += 2
                else:
                    j += 1
            else:
                i += 2
        else:
            i += 1
    return out


_BARE_DOLLAR_MATH_RE = re.compile(
    r"(?<!\$)\$(?!\$)((?:\\.|[^$\n])+?)\$(?!\$)"
)


def _add_rich_content_to_paragraph(
    paragraph,
    text: str,
    base_dir: Path | None,
    *,
    formula_inline_max_h_in: float = _FORMULA_INLINE_MAX_H_IN,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
    mono: bool = False,
) -> None:
    """同一段内混排文字（**粗体**/`代码`）与公式/插图（含 HTML 注释隐藏引用）。"""
    taken: list[tuple[int, int]] = []
    tokens: list[tuple[int, int, str, tuple]] = []

    for m in _INLINE_MATH_WITH_HIDDEN_IMG_RE.finditer(text):
        tokens.append(
            (
                m.start(),
                m.end(),
                "math_omml_or_img",
                (m.group(1), m.group(2), m.group(3).strip()),
            )
        )
        taken.append((m.start(), m.end()))

    for start, end, latex, alt, src in _iter_inline_paren_math_with_img(text):
        if _span_overlaps(taken, start, end):
            continue
        tokens.append((start, end, "math_omml_or_img", (latex, alt, src)))
        taken.append((start, end))

    for m in _HIDDEN_MD_IMAGE_COMMENT_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "hidden_img", (m.group(1), m.group(2).strip())))
        taken.append((m.start(), m.end()))

    for m in _MD_IMAGE_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "visible_img", (m.group(1), m.group(2).strip())))
        taken.append((m.start(), m.end()))

    for start, end, latex in _iter_inline_paren_math_bare(text):
        if _span_overlaps(taken, start, end):
            continue
        tokens.append((start, end, "math_omml", (latex,)))
        taken.append((start, end))

    for m in _BARE_DOLLAR_MATH_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "math_omml", (m.group(1),)))
        taken.append((m.start(), m.end()))

    inline_pat = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`)")
    for m in inline_pat.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "inline", (m.group(1),)))
        taken.append((m.start(), m.end()))

    tokens.sort(key=lambda t: t[0])
    pos = 0
    for start, end, kind, payload in tokens:
        if start > pos:
            _add_inline_to_paragraph(paragraph, text[pos:start], mono=mono)
        if kind == "inline":
            token = payload[0]
            if token.startswith("**"):
                run = paragraph.add_run(token[2:-2])
                _set_run_font(run, "宋体", 10.5, bold=True)
            else:
                run = paragraph.add_run(token[1:-1])
                _set_run_font(run, "Consolas", 9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif kind == "math_omml":
            latex = payload[0]
            if _try_append_omml(paragraph, latex, display=False):
                _MATH_STATS.record(latex, "omml")
            else:
                _add_inline_to_paragraph(paragraph, text[start:end], mono=mono)
                _MATH_STATS.record(latex, "text")
        elif kind == "math_omml_or_img":
            latex, alt, src = payload[0], payload[1], payload[2]
            if _try_append_omml(paragraph, latex, display=False):
                _MATH_STATS.record(latex, "omml")
            else:
                _embed_from_image_ref(
                    alt,
                    src,
                    base_dir,
                    paragraph=paragraph,
                    image_max_w_in=image_max_w_in,
                    image_max_h_in=image_max_h_in,
                )
                _MATH_STATS.record(latex, "png")
        else:
            alt, src = payload[0], payload[1]
            _embed_from_image_ref(
                alt,
                src,
                base_dir,
                paragraph=paragraph,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
            )
        pos = end
    if pos < len(text):
        _add_inline_to_paragraph(paragraph, text[pos:], mono=mono)


def _set_run_font(run, name: str = "宋体", size_pt: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def _add_inline_to_paragraph(paragraph, text: str, *, mono: bool = False):
    """解析 **粗体**、`行内代码` 与普通文本，写入同一段落。"""
    if not text:
        return
    # 拆分为：粗体、行内代码、普通
    pattern = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            _set_run_font(run, "Consolas" if mono else "宋体", 10.5 if not mono else 9)
        token = m.group(1)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, "宋体", 10.5, bold=True)
        else:  # `code`
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, "Consolas", 9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, "Consolas" if mono else "宋体", 10.5 if not mono else 9)


def _add_heading(doc: Document, level: int, text: str):
    """level 1–9 对应 Word 标题 1–标题 9；去除行内标记时保留可读文本。"""
    plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    plain = re.sub(r"`([^`]+)`", r"\1", plain)
    h = doc.add_heading(plain, level=min(max(level, 1), 9))
    for run in h.runs:
        _set_run_font(run, "黑体" if level <= 2 else "宋体")


def _add_body_paragraph(
    doc: Document,
    text: str,
    base_dir: Path | None = None,
    *,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if _line_has_embeddable_images(text):
        _add_rich_content_to_paragraph(
            p,
            text,
            base_dir,
            image_max_w_in=_DEFAULT_IMAGE_MAX_W_IN,
            image_max_h_in=image_max_h_in,
        )
    else:
        _add_inline_to_paragraph(p, text)
    for run in p.runs:
        if run.font.name in (None, ""):
            _set_run_font(run, "宋体", 10.5)


def _add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    body = "\n".join(lines)
    run = p.add_run(body)
    _set_run_font(run, "Consolas", 9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def _paragraph_num_id(paragraph) -> int | None:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    el = num_pr.find(qn("w:numId"))
    if el is None:
        return None
    raw = el.get(qn("w:val"))
    if raw is None or not str(raw).isdigit():
        return None
    return int(raw)


def _style_num_id(doc: Document, style_name: str) -> int | None:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    p_pr = style._element.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    el = num_pr.find(qn("w:numId"))
    if el is None:
        return None
    raw = el.get(qn("w:val"))
    if raw is None or not str(raw).isdigit():
        return None
    return int(raw)


def _set_paragraph_num_id(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), "0")
        num_pr.append(ilvl_el)
        el = OxmlElement("w:numId")
        num_pr.append(el)
        p_pr.append(num_pr)
    else:
        el = num_pr.find(qn("w:numId"))
        if el is None:
            el = OxmlElement("w:numId")
            num_pr.append(el)
    el.set(qn("w:val"), str(num_id))


def _new_list_num_id(doc: Document, src_num_id: int) -> int:
    """克隆一份编号实例，使新一组有序列表从 1 起计。"""
    numbering_part = getattr(doc.part, "numbering_part", None)
    if numbering_part is None:
        return src_num_id
    numbering_elm = numbering_part._element
    src = None
    used = [0]
    for node in numbering_elm.findall(qn("w:num")):
        raw = node.get(qn("w:numId"))
        if raw and str(raw).isdigit():
            nid = int(raw)
            used.append(nid)
            if nid == src_num_id:
                src = node
    if src is None:
        return src_num_id
    abs_el = src.find(qn("w:abstractNumId"))
    if abs_el is None:
        return src_num_id
    new_id = max(used) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_id))
    new_abs = OxmlElement("w:abstractNumId")
    new_abs.set(qn("w:val"), abs_el.get(qn("w:val")))
    new_num.append(new_abs)
    lvl_ov = OxmlElement("w:lvlOverride")
    lvl_ov.set(qn("w:ilvl"), "0")
    start_ov = OxmlElement("w:startOverride")
    start_ov.set(qn("w:val"), "1")
    lvl_ov.append(start_ov)
    new_num.append(lvl_ov)
    numbering_elm.append(new_num)
    return new_id


def _add_list_item(
    doc: Document,
    text: str,
    ordered: bool,
    base_dir: Path | None,
    *,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
):
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except (KeyError, ValueError):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(3)
    if _line_has_embeddable_images(text):
        _add_rich_content_to_paragraph(
            p,
            text,
            base_dir,
            image_max_w_in=_DEFAULT_IMAGE_MAX_W_IN,
            image_max_h_in=image_max_h_in,
        )
    else:
        _add_inline_to_paragraph(p, text)
    for run in p.runs:
        _set_run_font(run, "宋体", 10.5)
    return p


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _split_table_cells(line: str) -> list[str]:
    """按列分隔符 ``|`` 拆分表格行，忽略 ``\\(...\\)``、``$...$``、``<!-- -->`` 与 ``\\|`` 内的竖线。"""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]

    cells: list[str] = []
    buf: list[str] = []
    i = 0
    n = len(s)

    while i < n:
        if s.startswith("<!--", i):
            end = s.find("-->", i)
            if end == -1:
                buf.append(s[i:])
                break
            buf.append(s[i : end + 3])
            i = end + 3
            continue

        if s.startswith("\\(", i):
            end = s.find("\\)", i + 2)
            if end == -1:
                buf.append(s[i:])
                break
            buf.append(s[i : end + 2])
            i = end + 2
            continue

        if s[i] == "$":
            if i + 1 < n and s[i + 1] == "$":
                end = s.find("$$", i + 2)
                if end == -1:
                    buf.append(s[i:])
                    break
                buf.append(s[i : end + 2])
                i = end + 2
                continue
            j = i + 1
            while j < n:
                if s[j] == "$" and (j == 0 or s[j - 1] != "\\"):
                    buf.append(s[i : j + 1])
                    i = j + 1
                    break
                j += 1
            else:
                buf.append(s[i:])
                break
            continue

        if s[i] == "\\" and i + 1 < n and s[i + 1] == "|":
            buf.append("\\|")
            i += 2
            continue

        if s[i] == "|":
            cells.append("".join(buf).strip())
            buf = []
            i += 1
            continue

        buf.append(s[i])
        i += 1

    cells.append("".join(buf).strip())
    return cells


def _parse_table_row(line: str) -> list[str]:
    return _split_table_cells(line)


def _is_table_sep(row: list[str]) -> bool:
    if not row:
        return False
    return all(re.match(r"^:?-{3,}:?$", c.strip()) for c in row if c.strip())


def _add_table(doc: Document, rows: list[list[str]], base_dir: Path | None = None):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            if _line_has_embeddable_images(cell_text):
                _add_rich_content_to_paragraph(p, cell_text, base_dir)
            else:
                _add_inline_to_paragraph(p, cell_text)
            for run in p.runs:
                _set_run_font(run, "宋体", 10)


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("─" * 32)
    _set_run_font(run, "宋体", 8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def _resolve_image_path(src: str, base_dir: Path | None) -> Path | None:
    if not base_dir:
        return None
    path = (base_dir / src).resolve() if not Path(src).is_absolute() else Path(src)
    return path if path.is_file() else None


def _try_add_image(
    doc: Document,
    line: str,
    base_dir: Path | None,
    *,
    max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> bool:
    m = _MD_IMAGE_RE.match(line.strip())
    if not m or not base_dir:
        return False
    alt, src = m.group(1), m.group(2).strip()
    _embed_from_image_ref(
        alt,
        src,
        base_dir,
        doc=doc,
        image_max_w_in=max_w_in,
        image_max_h_in=max_h_in,
    )
    return True


def _line_has_embeddable_images(line: str) -> bool:
    if (
        _MD_IMAGE_RE.search(line)
        or _HIDDEN_MD_IMAGE_COMMENT_RE.search(line)
        or _INLINE_MATH_WITH_HIDDEN_IMG_RE.search(line)
        or _BARE_DOLLAR_MATH_RE.search(line)
    ):
        return True
    if _iter_inline_paren_math_with_img(line) or _iter_inline_paren_math_bare(line):
        return True
    return False


def _add_paragraph_with_inline_images(
    doc: Document,
    line: str,
    base_dir: Path | None,
    *,
    max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> None:
    """段落内混排文字与公式/插图（含 HTML 注释隐藏引用）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    _add_rich_content_to_paragraph(
        p,
        line,
        base_dir,
        image_max_w_in=max_w_in,
        image_max_h_in=max_h_in,
    )
    for run in p.runs:
        if run.font.name in (None, ""):
            _set_run_font(run, "宋体", 10.5)


def convert_md_to_docx(
    md_text: str,
    base_dir: Path | None,
    *,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
    prefer_omml: bool = True,
) -> Document:
    global _PREFER_OMML
    _PREFER_OMML = bool(prefer_omml)
    _MATH_STATS.reset()
    doc = Document()
    # 默认正文样式
    try:
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
    except (AttributeError, KeyError):
        pass

    lines = md_text.splitlines()
    i = 0
    para_buf: list[str] = []
    ordered_num_id: int | None = None

    def break_ordered_list() -> None:
        nonlocal ordered_num_id
        ordered_num_id = None

    def bind_ordered_paragraph(paragraph) -> None:
        nonlocal ordered_num_id
        src = _paragraph_num_id(paragraph) or _style_num_id(doc, "List Number")
        if src is None:
            return
        if ordered_num_id is None:
            ordered_num_id = _new_list_num_id(doc, src)
        _set_paragraph_num_id(paragraph, ordered_num_id)

    def flush_paragraph():
        nonlocal para_buf
        if not para_buf:
            return
        break_ordered_list()
        # 每行独立成段，避免「（1）…\n（2）…」被空格拼成一段（Word 内不换行）
        for p in para_buf:
            t = p.strip()
            if t:
                _add_body_paragraph(
                    doc,
                    t,
                    base_dir,
                    image_max_h_in=image_max_h_in,
                )
        para_buf = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        if line.strip() == "":
            flush_paragraph()
            i += 1
            continue

        # 围栏代码块
        if line.strip().startswith("```"):
            flush_paragraph()
            break_ordered_list()
            fence_lang = line.strip()[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            # 定稿 MD 保留 mermaid 源码 + 图示注释：Word 只嵌 PNG，不写源码块
            if fence_lang.lower() == "mermaid":
                j = i
                while j < len(lines) and lines[j].strip() == "":
                    j += 1
                if j < len(lines):
                    cm = _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[j].strip())
                    if cm and _is_diagram_image(cm.group(1), cm.group(2).strip()):
                        continue
            _add_code_block(doc, code_lines)
            continue

        # 块级公式：\[ ... \] + 可选 HTML 注释（OMML → PNG → 原文）
        if line.strip() == "\\[":
            flush_paragraph()
            break_ordered_list()
            i += 1
            math_lines: list[str] = []
            while i < len(lines) and lines[i].strip() != "\\]":
                math_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            hidden: tuple[str, str] | None = None
            if i < len(lines):
                cm = _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[i].strip())
                if cm:
                    hidden = (cm.group(1), cm.group(2).strip())
                    i += 1
            latex = "\n".join(ln.rstrip("\n") for ln in math_lines).strip()
            _add_block_equation(
                doc,
                latex,
                base_dir=base_dir,
                hidden=hidden,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
                raw_fallback_lines=["\\[", *math_lines, "\\]"],
            )
            continue

        # 块级公式：$$ ... $$ + 可选 HTML 注释（OMML → PNG → 原文）
        if line.strip() == "$$":
            flush_paragraph()
            break_ordered_list()
            i += 1
            math_lines = []
            while i < len(lines) and lines[i].strip() != "$$":
                math_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            hidden = None
            if i < len(lines):
                cm = _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[i].strip())
                if cm:
                    hidden = (cm.group(1), cm.group(2).strip())
                    i += 1
            latex = "\n".join(ln.rstrip("\n") for ln in math_lines).strip()
            _add_block_equation(
                doc,
                latex,
                base_dir=base_dir,
                hidden=hidden,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
                raw_fallback_lines=math_lines,
            )
            continue

        # 单行 $$ ... $$
        strip = line.strip()
        if strip.startswith("$$") and strip.endswith("$$") and len(strip) > 4:
            flush_paragraph()
            break_ordered_list()
            latex = strip[2:-2].strip()
            i += 1
            hidden = None
            if i < len(lines):
                cm = _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[i].strip())
                if cm:
                    hidden = (cm.group(1), cm.group(2).strip())
                    i += 1
            _add_block_equation(
                doc,
                latex,
                base_dir=base_dir,
                hidden=hidden,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
                raw_fallback_lines=[strip],
            )
            continue

        # 独立 HTML 注释行（公式图 / mermaid 框图引用）
        if _HIDDEN_MD_IMAGE_COMMENT_RE.fullmatch(line.strip()):
            flush_paragraph()
            break_ordered_list()
            _try_embed_hidden_comment_line(
                doc,
                line,
                base_dir,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
            )
            i += 1
            continue

        # 图片行或含行内公式/注释的段落
        if _line_has_embeddable_images(line):
            flush_paragraph()
            break_ordered_list()
            stripped = line.strip()
            if _MD_IMAGE_RE.fullmatch(stripped) or (
                stripped.startswith("![") and stripped.count("![") == 1
            ):
                _try_add_image(
                    doc,
                    line,
                    base_dir,
                    max_w_in=image_max_w_in,
                    max_h_in=image_max_h_in,
                )
            else:
                _add_paragraph_with_inline_images(
                    doc,
                    line,
                    base_dir,
                    max_w_in=image_max_w_in,
                    max_h_in=image_max_h_in,
                )
            i += 1
            continue

        # 水平线
        if re.match(r"^[\s\-*_]{3,}\s*$", line) and set(line.strip()) <= {"-", "*", "_", " "}:
            flush_paragraph()
            break_ordered_list()
            _add_horizontal_rule(doc)
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            flush_paragraph()
            break_ordered_list()
            level = len(m.group(1))
            title = m.group(2).strip()
            title = re.sub(r"\s+#+\s*$", "", title)
            _add_heading(doc, level, title)
            i += 1
            continue

        # 引用
        if line.lstrip().startswith("> "):
            flush_paragraph()
            break_ordered_list()
            quote = line.lstrip()[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            _add_inline_to_paragraph(p, quote)
            for run in p.runs:
                _set_run_font(run, "宋体", 10.5)
            i += 1
            continue

        # 表格块
        if _is_table_row(line):
            flush_paragraph()
            break_ordered_list()
            table_rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                row = _parse_table_row(lines[i])
                if not _is_table_sep(row):
                    table_rows.append(row)
                i += 1
            _add_table(doc, table_rows, base_dir)
            continue

        # 无序列表
        um = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if um:
            flush_paragraph()
            break_ordered_list()
            _add_list_item(
                doc,
                um.group(2).strip(),
                ordered=False,
                base_dir=base_dir,
                image_max_h_in=image_max_h_in,
            )
            i += 1
            continue

        # 有序列表
        om = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if om:
            flush_paragraph()
            p = _add_list_item(
                doc,
                om.group(2).strip(),
                ordered=True,
                base_dir=base_dir,
                image_max_h_in=image_max_h_in,
            )
            bind_ordered_paragraph(p)
            i += 1
            continue

        para_buf.append(line)
        i += 1

    flush_paragraph()
    return doc


def main(argv: list[str] | None = None) -> int:
    ensure_utf8_stdio()
    p = argparse.ArgumentParser(description="Markdown → Word（标题样式映射）")
    p.add_argument("-i", "--input", required=True, help="输入 .md 路径")
    p.add_argument("-o", "--output", required=True, help="输出 .docx 路径")
    p.add_argument(
        "--base-dir",
        default=None,
        help="解析 ![](/相对路径) 图片时的根目录（默认使用 .md 所在目录）",
    )
    p.add_argument(
        "--image-max-width-inches",
        type=float,
        default=_DEFAULT_IMAGE_MAX_W_IN,
        metavar="IN",
        help=f"插图最大宽度（英寸，默认 {_DEFAULT_IMAGE_MAX_W_IN}），与高度共同约束等比缩放",
    )
    p.add_argument(
        "--image-max-height-inches",
        type=float,
        default=_DEFAULT_IMAGE_MAX_H_IN,
        metavar="IN",
        help=f"插图最大高度（英寸，默认 {_DEFAULT_IMAGE_MAX_H_IN}），避免竖图仅按宽度缩放后超出单页可视区域",
    )
    p.add_argument(
        "--math-render",
        action="store_true",
        help="预渲染公式 PNG 供 OMML 失败时嵌入（须 matplotlib；默认跳过）",
    )
    p.add_argument(
        "--no-math-render",
        action="store_true",
        help=argparse.SUPPRESS,  # 旧开关；现默认已跳过公式 PNG
    )
    p.add_argument(
        "--no-omml",
        action="store_true",
        help="不写入可编辑 Office Math，仅用 PNG/原文（旧行为）",
    )
    args = p.parse_args(argv)

    in_path = Path(args.input).resolve()
    if not in_path.is_file():
        print(f"错误：找不到输入文件 {in_path}", file=sys.stderr)
        return 1

    base = Path(args.base_dir).resolve() if args.base_dir else in_path.parent
    try:
        md_text = in_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        md_text = in_path.read_text(encoding="utf-8", errors="replace")
        print("警告：输入文件含非 UTF-8 字节，已使用替换字符解码后继续转换。", file=sys.stderr)

    if args.math_render:
        md_text = _maybe_render_math_md(md_text, base)

    doc = convert_md_to_docx(
        md_text,
        base_dir=base,
        image_max_w_in=args.image_max_width_inches,
        image_max_h_in=args.image_max_height_inches,
        prefer_omml=not args.no_omml,
    )
    out_path = Path(args.output).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"DOCX: ok=1", file=sys.stderr)
    print(f"已写入: {out_path}")
    _MATH_STATS.report()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
