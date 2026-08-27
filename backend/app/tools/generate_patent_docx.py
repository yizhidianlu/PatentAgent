#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate a Chinese patent application DOCX from structured JSON.

The script intentionally uses only the Python standard library so the
paper2patent skill can run in repositories without a package manager. It also
supports embedding black-and-white SVG patent drawings referenced by the
`drawing_assets` field.
"""

from __future__ import annotations


import argparse
import json
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

try:  # Optional, used when available to create a more Word-compatible package.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    PYTHON_DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - depends on local environment
    PYTHON_DOCX_AVAILABLE = False


A4_WIDTH_DXA = 11906
A4_HEIGHT_DXA = 16838
MARGIN_DXA = 1440
EMU_PER_INCH = 914400
DRAWING_WIDTH_EMU = int(5.8 * EMU_PER_INCH)
MAX_DRAWING_HEIGHT_EMU = int(4.7 * EMU_PER_INCH)


def read_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8-sig") as fh:
        data = json.load(fh)
    if not isinstance(data, dict):
        raise ValueError("Input JSON must be an object.")
    return data


def as_paragraphs(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        result: list[str] = []
        for item in value:
            result.extend(as_paragraphs(item))
        return result
    if isinstance(value, dict):
        result = []
        for _, item in value.items():
            result.extend(as_paragraphs(item))
        return result
    text = str(value).strip()
    if not text:
        return []
    return [part.strip() for part in re.split(r"(?:\r?\n){1,}", text) if part.strip()]


def ensure_list(value: Any, placeholder: str) -> list[str]:
    items = as_paragraphs(value)
    return items or [placeholder]


def xml_text(text: str) -> str:
    return escape(text, {'"': "&quot;"})


def run(text: str, bold: bool = False, size: int | None = None) -> str:
    props = []
    if bold:
        props.append("<w:b/>")
    if size is not None:
        props.append(f'<w:sz w:val="{size}"/>')
        props.append(f'<w:szCs w:val="{size}"/>')
    props_xml = f"<w:rPr>{''.join(props)}</w:rPr>" if props else ""
    return f'<w:r>{props_xml}<w:t xml:space="preserve">{xml_text(text)}</w:t></w:r>'


def paragraph(
    text: str,
    style: str | None = None,
    align: str | None = None,
    bold: bool = False,
    size: int | None = None,
) -> str:
    p_props = []
    if style:
        p_props.append(f'<w:pStyle w:val="{style}"/>')
    if align:
        p_props.append(f'<w:jc w:val="{align}"/>')
    p_pr = f"<w:pPr>{''.join(p_props)}</w:pPr>" if p_props else ""
    return f"<w:p>{p_pr}{run(text, bold=bold, size=size)}</w:p>"


def section_heading(text: str) -> str:
    return paragraph(text, style="Heading1", bold=True, size=32)


def subsection_heading(text: str) -> str:
    return paragraph(text, style="Heading2", bold=True, size=28)


def caption_paragraph(text: str) -> str:
    return paragraph(text, style="Caption", align="center", size=21)


def body_paragraphs(values: list[str]) -> list[str]:
    return [paragraph(value, style="PatentBody") for value in values]


def description_blocks(description: Any) -> list[str]:
    if isinstance(description, dict):
        mapping = [
            ("technical_field", "技术领域"),
            ("background", "背景技术"),
            ("invention_content", "发明内容"),
            ("drawing_description", "附图说明"),
            ("embodiments", "具体实施方式"),
        ]
        blocks: list[str] = []
        for key, heading in mapping:
            values = as_paragraphs(description.get(key))
            if not values:
                values = [f"【待补充：{heading}。】"]
            blocks.append(subsection_heading(heading))
            blocks.extend(body_paragraphs(values))
        return blocks
    return body_paragraphs(ensure_list(description, "【待补充：说明书正文。】"))


def resolve_asset_path(raw_path: str, base_dir: Path) -> Path:
    path = Path(raw_path)
    if path.is_absolute():
        return path
    base_candidate = base_dir / path
    if base_candidate.exists():
        return base_candidate.resolve()
    cwd_candidate = Path.cwd() / path
    if cwd_candidate.exists():
        return cwd_candidate.resolve()
    return base_candidate.resolve()


def svg_dimensions(path: Path) -> tuple[float, float]:
    text = path.read_text(encoding="utf-8", errors="ignore")[:4096]
    tag_match = re.search(r"<svg\b[^>]*>", text, flags=re.IGNORECASE)
    tag = tag_match.group(0) if tag_match else text
    width_match = re.search(r'\bwidth="([0-9.]+)', tag)
    height_match = re.search(r'\bheight="([0-9.]+)', tag)
    if width_match and height_match:
        width = float(width_match.group(1))
        height = float(height_match.group(1))
        if width > 0 and height > 0:
            return width, height
    viewbox_match = re.search(
        r'\bviewBox="[-0-9.]+\s+[-0-9.]+\s+([0-9.]+)\s+([0-9.]+)"', tag
    )
    if viewbox_match:
        width = float(viewbox_match.group(1))
        height = float(viewbox_match.group(2))
        if width > 0 and height > 0:
            return width, height
    return 1000.0, 600.0


def drawing_extent(path: Path) -> tuple[int, int]:
    width, height = svg_dimensions(path)
    cx = DRAWING_WIDTH_EMU
    cy = int(cx * height / width)
    if cy > MAX_DRAWING_HEIGHT_EMU:
        cy = MAX_DRAWING_HEIGHT_EMU
        cx = int(cy * width / height)
    return cx, cy


def collect_drawing_images(
    data: dict[str, Any], base_dir: Path, require_drawings: bool
) -> list[dict[str, Any]]:
    assets = data.get("drawing_assets")
    if not isinstance(assets, list):
        assets = []

    images: list[dict[str, Any]] = []
    next_rid = 2
    for index, asset in enumerate(assets, start=1):
        if not isinstance(asset, dict):
            continue
        raw_path = str(asset.get("svg_path") or "").strip()
        if not raw_path:
            continue
        path = resolve_asset_path(raw_path, base_dir)
        if not path.exists():
            raise FileNotFoundError(f"Drawing asset does not exist: {path}")
        if path.suffix.lower() != ".svg":
            raise ValueError(f"Only SVG drawing assets are supported: {path}")
        png_path: Path | None = None
        raw_png_path = str(asset.get("png_path") or "").strip()
        if raw_png_path:
            png_path = resolve_asset_path(raw_png_path, base_dir)
            if not png_path.exists():
                raise FileNotFoundError(f"PNG fallback asset does not exist: {png_path}")
            if png_path.suffix.lower() != ".png":
                raise ValueError(f"Only PNG fallback assets are supported: {png_path}")
        cx, cy = drawing_extent(path)
        figure_no = int(asset.get("figure_no") or index)
        title = str(asset.get("title") or f"图{figure_no}").strip()
        caption = str(asset.get("caption") or f"图{figure_no} {title}").strip()
        primary_path = png_path or path
        primary_ext = primary_path.suffix.lower().lstrip(".")
        primary_rid = f"rId{next_rid}"
        next_rid += 1
        svg_rid = None
        svg_target = None
        if png_path:
            svg_rid = f"rId{next_rid}"
            svg_target = f"media/figure_{figure_no}.svg"
            next_rid += 1
        images.append(
            {
                "asset": asset,
                "path": primary_path,
                "rid": primary_rid,
                "target": f"media/figure_{figure_no}.{primary_ext}",
                "svg_path": path,
                "svg_rid": svg_rid,
                "svg_target": svg_target,
                "figure_no": figure_no,
                "title": title,
                "caption": caption,
                "cx": cx,
                "cy": cy,
            }
        )

    images.sort(key=lambda item: item["figure_no"])
    if require_drawings and not images:
        raise RuntimeError(
            "No drawing_assets with SVG files were found. Run generate_patent_drawings.py first."
        )
    return images


def image_paragraph(image: dict[str, Any], doc_pr_id: int) -> str:
    cx = int(image["cx"])
    cy = int(image["cy"])
    rid = str(image["rid"])
    name = xml_text(str(image["caption"]))
    if image.get("svg_rid"):
        blip = (
            f'<a:blip r:embed="{rid}">'
            '<a:extLst>'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            f'<asvg:svgBlip r:embed="{image["svg_rid"]}"/>'
            '</a:ext>'
            '</a:extLst>'
            '</a:blip>'
        )
    else:
        blip = f'<a:blip r:embed="{rid}"/>'
    return f"""<w:p>
  <w:pPr><w:jc w:val="center"/><w:spacing w:before="160" w:after="80"/></w:pPr>
  <w:r>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="{doc_pr_id}" name="{name}"/>
        <wp:cNvGraphicFramePr>
          <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic>
              <pic:nvPicPr>
                <pic:cNvPr id="{doc_pr_id}" name="{name}"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                {blip}
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{cx}" cy="{cy}"/>
                </a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>"""


def select_abstract_image(data: dict[str, Any], images: list[dict[str, Any]]) -> dict[str, Any] | None:
    if not images:
        return None
    abstract_text = str(data.get("abstract_drawing") or "")
    match = re.search(r"图\s*(\d+)", abstract_text)
    if match:
        figure_no = int(match.group(1))
        for image in images:
            if image["figure_no"] == figure_no:
                return image
    for image in images:
        if bool(image["asset"].get("abstract_candidate")):
            return image
    return images[0]


def build_document_xml(data: dict[str, Any], images: list[dict[str, Any]]) -> str:
    invention_name = str(
        data.get("invention_name") or data.get("title") or "【待补充：发明名称】"
    ).strip()
    source_title = str(data.get("source_title") or "").strip()

    parts: list[str] = [
        paragraph(invention_name, style="Title", align="center", bold=True, size=38)
    ]
    if source_title:
        parts.append(paragraph(f"来源论文：{source_title}", align="center", size=20))

    parts.append(section_heading("一、说明书摘要"))
    parts.extend(body_paragraphs(ensure_list(data.get("abstract"), "【待补充：说明书摘要。】")))

    parts.append(section_heading("二、摘要附图"))
    parts.extend(
        body_paragraphs(
            ensure_list(data.get("abstract_drawing"), "【待补充：摘要附图选择或说明。】")
        )
    )
    next_doc_pr_id = 1
    abstract_image = select_abstract_image(data, images)
    if abstract_image:
        parts.append(image_paragraph(abstract_image, next_doc_pr_id))
        next_doc_pr_id += 1
        parts.append(caption_paragraph(str(abstract_image["caption"])))

    parts.append(section_heading("三、权利要求书"))
    parts.extend(body_paragraphs(ensure_list(data.get("claims"), "【待补充：权利要求书。】")))

    parts.append(section_heading("四、说明书"))
    parts.extend(description_blocks(data.get("description")))

    parts.append(section_heading("五、说明书附图"))
    if images:
        for image in images:
            parts.append(image_paragraph(image, next_doc_pr_id))
            next_doc_pr_id += 1
            parts.append(caption_paragraph(str(image["caption"])))
            spec = str(image["asset"].get("spec") or "").strip()
            if spec:
                parts.extend(body_paragraphs(as_paragraphs(spec)))
    else:
        parts.extend(body_paragraphs(ensure_list(data.get("drawings"), "【待补充：说明书附图。】")))

    gaps = as_paragraphs(data.get("gaps"))
    if gaps:
        parts.append(section_heading("材料缺口说明"))
        parts.extend(body_paragraphs(gaps))

    sect_pr = (
        "<w:sectPr>"
        f'<w:pgSz w:w="{A4_WIDTH_DXA}" w:h="{A4_HEIGHT_DXA}"/>'
        f'<w:pgMar w:top="{MARGIN_DXA}" w:right="{MARGIN_DXA}" '
        f'w:bottom="{MARGIN_DXA}" w:left="{MARGIN_DXA}" '
        'w:header="720" w:footer="720" w:gutter="0"/>'
        "</w:sectPr>"
    )
    body = "".join(parts) + sect_pr
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f"<w:body>{body}</w:body>"
        "</w:document>"
    )


def styles_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr>
      <w:rFonts w:ascii="SimSun" w:eastAsia="宋体" w:hAnsi="SimSun" w:cs="SimSun"/>
      <w:sz w:val="24"/>
      <w:szCs w:val="24"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Title">
    <w:name w:val="Title"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr>
      <w:spacing w:before="240" w:after="240"/>
      <w:jc w:val="center"/>
    </w:pPr>
    <w:rPr>
      <w:b/>
      <w:rFonts w:ascii="SimSun" w:eastAsia="宋体" w:hAnsi="SimSun" w:cs="SimSun"/>
      <w:sz w:val="38"/>
      <w:szCs w:val="38"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="Heading 1"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr>
      <w:spacing w:before="260" w:after="140"/>
      <w:keepNext/>
      <w:outlineLvl w:val="0"/>
    </w:pPr>
    <w:rPr>
      <w:b/>
      <w:rFonts w:ascii="SimSun" w:eastAsia="宋体" w:hAnsi="SimSun" w:cs="SimSun"/>
      <w:sz w:val="32"/>
      <w:szCs w:val="32"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="Heading 2"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr>
      <w:spacing w:before="180" w:after="80"/>
      <w:keepNext/>
      <w:outlineLvl w:val="1"/>
    </w:pPr>
    <w:rPr>
      <w:b/>
      <w:rFonts w:ascii="SimSun" w:eastAsia="宋体" w:hAnsi="SimSun" w:cs="SimSun"/>
      <w:sz w:val="28"/>
      <w:szCs w:val="28"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="PatentBody">
    <w:name w:val="Patent Body"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr>
      <w:spacing w:line="360" w:lineRule="auto" w:after="90"/>
      <w:ind w:firstLine="480"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="SimSun" w:eastAsia="宋体" w:hAnsi="SimSun" w:cs="SimSun"/>
      <w:sz w:val="24"/>
      <w:szCs w:val="24"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Caption">
    <w:name w:val="Caption"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr>
      <w:spacing w:before="60" w:after="160"/>
      <w:jc w:val="center"/>
      <w:keepNext/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="SimSun" w:eastAsia="宋体" w:hAnsi="SimSun" w:cs="SimSun"/>
      <w:sz w:val="21"/>
      <w:szCs w:val="21"/>
    </w:rPr>
  </w:style>
</w:styles>
"""


def content_types_xml(images: list[dict[str, Any]]) -> str:
    has_svg = any(image["target"].endswith(".svg") or image.get("svg_target") for image in images)
    has_png = any(image["target"].endswith(".png") for image in images)
    image_defaults = ""
    if has_svg:
        image_defaults += '  <Default Extension="svg" ContentType="image/svg+xml"/>\n'
    if has_png:
        image_defaults += '  <Default Extension="png" ContentType="image/png"/>\n'
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
{image_defaults}  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""


def package_rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""


def document_rels_xml(images: list[dict[str, Any]]) -> str:
    relationships = [
        '  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
    ]
    for image in images:
        relationships.append(
            f'  <Relationship Id="{image["rid"]}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            f'Target="{xml_text(str(image["target"]))}"/>'
        )
        if image.get("svg_rid") and image.get("svg_target"):
            relationships.append(
                f'  <Relationship Id="{image["svg_rid"]}" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                f'Target="{xml_text(str(image["svg_target"]))}"/>'
            )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">\n'
        + "\n".join(relationships)
        + "\n</Relationships>\n"
    )


def set_east_asian_font(run: Any, size_pt: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "SimSun"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def configure_python_docx_styles(document: Any) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_python_docx_paragraph(
    document: Any,
    text: str,
    size_pt: int = 12,
    bold: bool = False,
    align: Any | None = None,
    first_line: bool = False,
    before_pt: int = 0,
    after_pt: int = 4,
) -> Any:
    paragraph_obj = document.add_paragraph()
    if align is not None:
        paragraph_obj.alignment = align
    paragraph_obj.paragraph_format.space_before = Pt(before_pt)
    paragraph_obj.paragraph_format.space_after = Pt(after_pt)
    if first_line:
        paragraph_obj.paragraph_format.first_line_indent = Pt(24)
        paragraph_obj.paragraph_format.line_spacing = 1.5
    run_obj = paragraph_obj.add_run(text)
    set_east_asian_font(run_obj, size_pt=size_pt, bold=bold)
    return paragraph_obj


def add_python_docx_section_heading(document: Any, text: str) -> None:
    add_python_docx_paragraph(
        document,
        text,
        size_pt=16,
        bold=True,
        before_pt=14,
        after_pt=8,
    )


def add_python_docx_subsection_heading(document: Any, text: str) -> None:
    add_python_docx_paragraph(
        document,
        text,
        size_pt=14,
        bold=True,
        before_pt=10,
        after_pt=5,
    )


def add_python_docx_body(document: Any, values: list[str]) -> None:
    for value in values:
        add_python_docx_paragraph(document, value, first_line=True)


def add_python_docx_image(document: Any, image: dict[str, Any]) -> None:
    paragraph_obj = document.add_paragraph()
    paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_obj = paragraph_obj.add_run()
    run_obj.add_picture(str(image["path"]), width=Inches(5.8))
    add_python_docx_paragraph(
        document,
        str(image["caption"]),
        size_pt=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after_pt=8,
    )


def add_python_docx_description(document: Any, description: Any) -> None:
    if not isinstance(description, dict):
        add_python_docx_body(document, ensure_list(description, "【待补充：说明书正文。】"))
        return
    mapping = [
        ("technical_field", "技术领域"),
        ("background", "背景技术"),
        ("invention_content", "发明内容"),
        ("drawing_description", "附图说明"),
        ("embodiments", "具体实施方式"),
    ]
    for key, heading in mapping:
        values = as_paragraphs(description.get(key)) or [f"【待补充：{heading}。】"]
        add_python_docx_subsection_heading(document, heading)
        add_python_docx_body(document, values)


def add_svg_fallbacks_to_docx(output: Path, inserted_images: list[dict[str, Any]]) -> None:
    if not inserted_images:
        return

    with zipfile.ZipFile(output, "r") as source:
        document_xml = source.read("word/document.xml").decode("utf-8")
        rels_xml = source.read("word/_rels/document.xml.rels").decode("utf-8")
        content_types = source.read("[Content_Types].xml").decode("utf-8")
        original_items = {name: source.read(name) for name in source.namelist()}

    blip_matches = list(re.finditer(r'(<a:blip\b(?=[^>]*\br:embed="rId\d+")[^>]*)/>', document_xml))
    if not blip_matches:
        return

    existing_rids = [int(value) for value in re.findall(r'Id="rId(\d+)"', rels_xml)]
    next_rid = max(existing_rids, default=1) + 1
    fallback_items: list[tuple[str, Path, str]] = []
    replacements: dict[tuple[int, int], str] = {}

    for index, image in enumerate(inserted_images):
        if index >= len(blip_matches):
            break
        svg_path = image.get("svg_path")
        if not svg_path:
            continue
        svg_path = Path(svg_path)
        if not svg_path.exists():
            continue
        svg_rid = f"rId{next_rid}"
        next_rid += 1
        target = f"media/svg_fallback_{index + 1}.svg"
        match = blip_matches[index]
        replacements[(match.start(), match.end())] = (
            match.group(1)
            + ">"
            + "<a:extLst>"
            + '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            + f'<asvg:svgBlip r:embed="{svg_rid}"/>'
            + "</a:ext>"
            + "</a:extLst>"
            + "</a:blip>"
        )
        fallback_items.append((target, svg_path, svg_rid))

    if not fallback_items:
        return

    if "xmlns:asvg=" not in document_xml:
        document_xml = document_xml.replace(
            "<w:document ",
            '<w:document xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" ',
            1,
        )
    ignorable_match = re.search(r'mc:Ignorable="([^"]*)"', document_xml)
    if ignorable_match and "asvg" not in ignorable_match.group(1).split():
        updated = f'mc:Ignorable="{ignorable_match.group(1)} asvg"'
        document_xml = (
            document_xml[: ignorable_match.start()]
            + updated
            + document_xml[ignorable_match.end() :]
        )

    rebuilt = []
    cursor = 0
    for (start, end), replacement in sorted(replacements.items()):
        rebuilt.append(document_xml[cursor:start])
        rebuilt.append(replacement)
        cursor = end
    rebuilt.append(document_xml[cursor:])
    document_xml = "".join(rebuilt)

    relationship_rows = []
    for target, _, svg_rid in fallback_items:
        relationship_rows.append(
            f'<Relationship Id="{svg_rid}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            f'Target="{target}"/>'
        )
    rels_xml = rels_xml.replace(
        "</Relationships>", "".join(relationship_rows) + "</Relationships>"
    )

    if 'Extension="svg"' not in content_types:
        content_types = content_types.replace(
            '<Default Extension="xml" ContentType="application/xml"/>',
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Default Extension="svg" ContentType="image/svg+xml"/>',
        )

    temp = output.with_suffix(output.suffix + ".tmp")
    with zipfile.ZipFile(temp, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
        for name, payload in original_items.items():
            if name in {"word/document.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"}:
                continue
            target_zip.writestr(name, payload)
        target_zip.writestr("word/document.xml", document_xml)
        target_zip.writestr("word/_rels/document.xml.rels", rels_xml)
        target_zip.writestr("[Content_Types].xml", content_types)
        for target, svg_path, _ in fallback_items:
            target_zip.write(svg_path, f"word/{target}")
    temp.replace(output)


def write_docx_with_python_docx(
    data: dict[str, Any],
    output: Path,
    images: list[dict[str, Any]],
    embed_svg_fallback: bool = False,
) -> int:
    document = Document()
    configure_python_docx_styles(document)
    invention_name = str(
        data.get("invention_name") or data.get("title") or "【待补充：发明名称】"
    ).strip()
    source_title = str(data.get("source_title") or "").strip()

    add_python_docx_paragraph(
        document,
        invention_name,
        size_pt=19,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before_pt=12,
        after_pt=12,
    )
    if source_title:
        add_python_docx_paragraph(
            document,
            f"来源论文：{source_title}",
            size_pt=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            after_pt=10,
        )

    add_python_docx_section_heading(document, "一、说明书摘要")
    add_python_docx_body(document, ensure_list(data.get("abstract"), "【待补充：说明书摘要。】"))

    inserted_images: list[dict[str, Any]] = []
    add_python_docx_section_heading(document, "二、摘要附图")
    add_python_docx_body(
        document, ensure_list(data.get("abstract_drawing"), "【待补充：摘要附图选择或说明。】")
    )
    abstract_image = select_abstract_image(data, images)
    if abstract_image:
        add_python_docx_image(document, abstract_image)
        inserted_images.append(abstract_image)

    add_python_docx_section_heading(document, "三、权利要求书")
    add_python_docx_body(document, ensure_list(data.get("claims"), "【待补充：权利要求书。】"))

    add_python_docx_section_heading(document, "四、说明书")
    add_python_docx_description(document, data.get("description"))

    add_python_docx_section_heading(document, "五、说明书附图")
    if images:
        for image in images:
            add_python_docx_image(document, image)
            inserted_images.append(image)
            spec = str(image["asset"].get("spec") or "").strip()
            if spec:
                add_python_docx_body(document, as_paragraphs(spec))
    else:
        add_python_docx_body(document, ensure_list(data.get("drawings"), "【待补充：说明书附图。】"))

    gaps = as_paragraphs(data.get("gaps"))
    if gaps:
        add_python_docx_section_heading(document, "材料缺口说明")
        add_python_docx_body(document, gaps)

    document.save(output)
    if embed_svg_fallback:
        add_svg_fallbacks_to_docx(output, inserted_images)
    return len(images)


def write_docx(
    data: dict[str, Any],
    output: Path,
    base_dir: Path,
    require_drawings: bool = False,
    embed_svg_fallback: bool = False,
) -> int:
    output.parent.mkdir(parents=True, exist_ok=True)
    images = collect_drawing_images(data, base_dir, require_drawings)
    if PYTHON_DOCX_AVAILABLE and all(Path(image["path"]).suffix.lower() == ".png" for image in images):
        return write_docx_with_python_docx(
            data, output, images, embed_svg_fallback=embed_svg_fallback
        )
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml(images))
        archive.writestr("_rels/.rels", package_rels_xml())
        archive.writestr("word/document.xml", build_document_xml(data, images))
        archive.writestr("word/styles.xml", styles_xml())
        archive.writestr("word/_rels/document.xml.rels", document_rels_xml(images))
        for image in images:
            archive.write(image["path"], f'word/{image["target"]}')
            if image.get("svg_rid") and image.get("svg_target"):
                archive.write(image["svg_path"], f'word/{image["svg_target"]}')
    return len(images)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a DOCX Chinese patent application from structured JSON."
    )
    parser.add_argument("input", type=Path, help="Path to patent content JSON.")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        required=True,
        help="Path to write the generated DOCX.",
    )
    parser.add_argument(
        "--require-drawings",
        action="store_true",
        help="Fail unless drawing_assets contains at least one embeddable SVG drawing.",
    )
    parser.add_argument(
        "--embed-svg-fallback",
        action="store_true",
        help="Add Office SVG fallback relationships in addition to PNG pictures. Use only when the target Word version supports SVG fallback markup.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    data = read_json(args.input)
    image_count = write_docx(
        data,
        args.output,
        base_dir=args.input.parent,
        require_drawings=args.require_drawings,
        embed_svg_fallback=args.embed_svg_fallback,
    )
    size = args.output.stat().st_size
    print(f"Wrote {args.output} ({size} bytes, embedded_drawings={image_count})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
