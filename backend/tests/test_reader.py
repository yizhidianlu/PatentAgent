"""模块 C（专利解读）端到端测试。

不依赖网络与真实 LLM：

- **样例专利**：用 pymupdf 现造一份中文专利 PDF（扉页著录项 + 权利要求书 +
  带四位段落号的说明书 + 附图说明），走真实的 `services/convert` PDF→md 管线；
- **取证**：monkeypatch `services/patent_fetch.fetch_patent_pdf` 返回这份本地 PDF；
  另有一例把它改成「抓取失败」，验证 `await_user` 门控与手工上传路径；
- **LLM**：脚本化 FakeLLM 按 step_key 回放预置产物，其余环节全部真跑 ——
  权要树 lint（脚本首版故意让权3 引用权4）、白话增量长度校验（脚本首版故意超 40 字）、
  说明书引用 lint（脚本正文故意写裸 `[0002]`、不存在的 `说明书 0099` 与半角连字符区间）、
  自检补丁（服务端确定性应用）、report.json 装配、reader_note_md 落盘与 docx 导出。
"""

from __future__ import annotations

import asyncio
import json
from pathlib import Path

from conftest import disk_path
from typing import Any

import pytest
from fastapi.testclient import TestClient

PUB_NO = "CN118888888A"
PATENT_TITLE = "一种超声图像病灶区域分割方法及系统"

# ---------------------------------------------------------------------------
# 一、样例专利 PDF（pymupdf 现造；china-ss 内置简体中文字体可原样抽回）
# ---------------------------------------------------------------------------

_FRONT_LINES = [
    "(19)国家知识产权局",
    "(12)发明专利申请",
    "(10)申请公布号 CN 118888888 A",
    "(43)申请公布日 2025-03-14",
    "(21)申请号 202410123456.7",
    "(22)申请日 2024-01-30",
    "(71)申请人 引途医疗科技（上海）有限公司",
    "(72)发明人 张三 李四",
    "(51)Int.Cl.",
    "G06T 7/11(2017.01)I",
    "A61B 8/08(2006.01)I",
    "(54)发明名称",
    PATENT_TITLE,
    "(57)摘要",
    "本发明公开一种超声图像病灶区域分割方法及系统，属于医学图像处理技术领域。",
]

CLAIM_1_TEXT = (
    "一种超声图像病灶区域分割方法，其特征在于，包括：对超声图像进行超像素划分，"
    "得到多个超像素节点；基于所述超像素节点构建邻接图；将所述邻接图输入图卷积网络，"
    "得到节点级病灶概率；根据所述节点级病灶概率生成病灶分割掩膜。"
)

_CLAIM_LINES = [
    "权利要求书",
    f"1.{CLAIM_1_TEXT}",
    "2.根据权利要求1所述的方法，其特征在于，所述超像素划分采用简单线性迭代聚类算法。",
    "3.根据权利要求1所述的方法，其特征在于，所述邻接图的节点属性包括灰度特征与纹理特征。",
    "4.根据权利要求2或3所述的方法，其特征在于，所述图卷积网络的层数为三层。",
    (
        "5.一种超声图像病灶区域分割系统，其特征在于，包括超像素划分模块、邻接图构建模块、"
        "特征聚合模块与掩膜生成模块。"
    ),
]

_DESC_LINES = [
    "说明书",
    PATENT_TITLE,
    "技术领域",
    "[0001]本发明属于医学图像处理技术领域，具体涉及一种超声图像病灶区域分割方法及系统。",
    "背景技术",
    "[0002]现有超声图像分割方法多采用逐像素分类网络，忽略相邻组织区域之间的上下文关系。",
    "[0003]由于斑点噪声普遍存在，逐像素方法的假阳性偏高，且计算冗余明显。",
    "发明内容",
    "[0004]本发明要解决的技术问题是提供一种能够显式表达区域间上下文关系的分割方法。",
    "[0005]为解决上述技术问题，本发明采用超像素划分与图卷积聚合相结合的技术方案。",
    "[0006]所述超像素划分是指按灰度相似性与空间邻近性将超声图像划分为若干超像素节点。",
    "[0007]所述邻接图是指以超像素节点为顶点、以空间相邻关系为边构成的无向图。",
    "附图说明",
    "[0008]图1为本发明实施例提供的超声图像病灶区域分割方法的流程图。",
    "[0009]图2为本发明实施例提供的超声图像病灶区域分割系统的结构示意图。",
    "具体实施方式",
    "[0010]下面结合附图对本发明的实施例进行详细说明。",
    "[0011]步骤S101，对超声图像进行超像素划分，得到多个超像素节点。",
    "[0012]步骤S102，基于超像素节点的灰度特征与纹理特征构建邻接图。",
    "[0013]步骤S103，将邻接图输入三层图卷积网络，逐层聚合邻域特征，输出节点级病灶概率。",
    "[0014]步骤S104，按节点级病灶概率归并超像素节点，生成病灶分割掩膜。",
    "[0015]在一个实施例中，超像素数量取值为200至400个。",
]


def _wrap(line: str, width: int = 34) -> list[str]:
    """按固定宽度硬折行（模拟真实专利 PDF 的排版换行）。"""
    if len(line) <= width:
        return [line]
    return [line[i : i + width] for i in range(0, len(line), width)]


def build_patent_pdf() -> bytes:
    """造一份可被 pymupdf 原样抽回文本的中文专利 PDF。"""
    import pymupdf

    physical: list[str] = []
    for block in (_FRONT_LINES, _CLAIM_LINES, _DESC_LINES):
        for logical in block:
            physical.extend(_wrap(logical))
        physical.append("")

    doc = pymupdf.open()
    page = None
    y = 0.0
    for line in physical:
        if page is None or y > 780:
            page = doc.new_page()
            y = 50.0
        if line:
            page.insert_text((45, y), line, fontname="china-ss", fontsize=10)
        y += 16
    data = doc.tobytes()
    doc.close()
    return data


PDF_BYTES = build_patent_pdf()


# ---------------------------------------------------------------------------
# 二、脚本化 LLM 产物
# ---------------------------------------------------------------------------


def _features(no: int, texts: list[str]) -> list[dict[str, str]]:
    return [{"fid": f"{no}-F{i}", "text": t} for i, t in enumerate(texts, 1)]


_CLAIM_NODES: list[dict[str, Any]] = [
    {
        "no": 1,
        "type": "independent",
        "refs": [],
        "ref_mode": "none",
        "preamble": "一种超声图像病灶区域分割方法，其特征在于，包括",
        "features": _features(
            1,
            [
                "对超声图像进行超像素划分，得到多个超像素节点",
                "基于所述超像素节点构建邻接图",
                "将所述邻接图输入图卷积网络，得到节点级病灶概率",
                "根据所述节点级病灶概率生成病灶分割掩膜",
            ],
        ),
    },
    {
        "no": 2,
        "type": "dependent",
        "refs": [1],
        "ref_mode": "single",
        "preamble": "根据权利要求1所述的方法，其特征在于",
        "features": _features(2, ["所述超像素划分采用简单线性迭代聚类算法"]),
    },
    {
        "no": 3,
        "type": "dependent",
        "refs": [1],
        "ref_mode": "single",
        "preamble": "根据权利要求1所述的方法，其特征在于",
        "features": _features(3, ["所述邻接图的节点属性包括灰度特征与纹理特征"]),
    },
    {
        "no": 4,
        "type": "dependent",
        "refs": [2, 3],
        "ref_mode": "alternative",
        "preamble": "根据权利要求2或3所述的方法，其特征在于",
        "features": _features(4, ["所述图卷积网络的层数为三层"]),
    },
    {
        "no": 5,
        "type": "independent",
        "refs": [],
        "ref_mode": "none",
        "preamble": "一种超声图像病灶区域分割系统，其特征在于，包括",
        "features": _features(
            5, ["超像素划分模块、邻接图构建模块、特征聚合模块与掩膜生成模块"]
        ),
    },
]

# 首版故意让权3 引用权4（引用号大于自身权号）——触发服务端树 lint + REPAIR
_BAD_CLAIM_NODES = [dict(node) for node in _CLAIM_NODES]
_BAD_CLAIM_NODES[2] = {**_CLAIM_NODES[2], "refs": [4]}

CLAIM_TREE_BAD: dict[str, Any] = {
    "claims": _BAD_CLAIM_NODES,
    "ambiguities": [
        {
            "claim": 4,
            "kind": "multi_ref",
            "issue": "权利要求 4 同时引用权 2 与权 3，需确认以哪一项为主引用",
            "candidates": [2, 3],
            "chosen": 2,
            "reason": "权 2 与权 3 分别细化不同环节，暂按择一引用处理",
        }
    ],
}

CLAIM_TREE_FIXED: dict[str, Any] = {
    "claims": _CLAIM_NODES,
    "ambiguities": CLAIM_TREE_BAD["ambiguities"],
}

# 首版权 4 的白话增量故意写到 40 字以上 → 触发长度校验与 REPAIR
DELTA_4_TOO_LONG = (
    "这一项在前面两项的基础上进一步把图卷积网络的层数固定成三层，"
    "既保证了邻域特征的聚合范围又避免了过深网络带来的过平滑问题"
)

CLAIM_DELTAS_BAD: dict[str, Any] = {
    "deltas": [
        {"claim": 1, "delta": "超像素划分＋邻接图＋图卷积，输出病灶分割掩膜"},
        {"claim": 2, "delta": "限定超像素划分用简单线性迭代聚类算法"},
        {"claim": 3, "delta": "限定节点属性为灰度特征与纹理特征两类"},
        {"claim": 4, "delta": DELTA_4_TOO_LONG},
        {"claim": 5, "delta": "同一方案的系统侧：四个功能模块串成处理链路"},
    ]
}

CLAIM_DELTAS_FIXED: dict[str, Any] = {
    "deltas": [
        {"claim": 1, "delta": "超像素划分＋邻接图＋图卷积，输出病灶分割掩膜"},
        {"claim": 2, "delta": "限定超像素划分用简单线性迭代聚类算法"},
        {"claim": 3, "delta": "限定节点属性为灰度特征与纹理特征两类"},
        {"claim": 4, "delta": "限定图卷积网络的层数固定为三层"},
        {"claim": 5, "delta": "同一方案的系统侧：四个功能模块串成处理链路"},
    ]
}

GROUNDING: dict[str, Any] = {
    "ipc_coordinates": [
        {"ipc_prefix": "A61", "industry": "生物医药与医疗器械", "basis": "hints_match"}
    ],
    "clues": [
        {
            "scene": "超声诊断工作站",
            "plain": "医生在超声设备上一键得到病灶轮廓，减少手工勾画。",
            "anchor_fits": [
                {
                    "kind": "feature",
                    "key": "F1",
                    "paragraph": "0011",
                    "fit": "该段落写了对超声图像做超像素划分，正是工作站里的预处理环节",
                }
            ],
            "confidence": "中",
            "disclaimer": "推断自公开文本，未联网核验",
        },
        {
            # 锚定到不存在的段落号 → 服务端应丢弃该线索并记 warning
            "scene": "凭空捏造的场景",
            "plain": "该线索没有任何专利内依据。",
            "anchor_fits": [
                {"kind": "claim", "key": "1", "paragraph": "9999", "fit": "不存在的段落"}
            ],
            "confidence": "低",
            "disclaimer": "推断自公开文本，未联网核验",
        },
    ],
    "no_clue_note": "",
}

STRUCTURE_MD = f"""## 三、权利要求树

共 5 项 · 独立 2 / 从属 3。独立权展开见第四节。

## 四、独立权利要求精读

> [!patent-claim] 权利要求 1
> 【{PUB_NO}·权利要求1】{CLAIM_1_TEXT}

| 特征 | 大白话 | 说明书依据 |
| --- | --- | --- |
| F1 | 把超声图像切成一块块超像素 | 说明书 0006 |
| F2 | 按相邻关系把超像素连成一张图 | 说明书 0007 |
| F3 | 图卷积算出每块是不是病灶 | 说明书 0013 |
| F4 | 把高概率的块拼成分割掩膜 | 说明书 0014 |

> [!patent-claim] 权利要求 5
> 【{PUB_NO}·权利要求5】一种超声图像病灶区域分割系统，其特征在于，包括超像素划分模块、邻接图构建模块、特征聚合模块与掩膜生成模块。

| 特征 | 大白话 | 说明书依据 |
| --- | --- | --- |
| F1 | 四个模块对应方法的四个步骤 | 说明书 0011-0014 |

## 五、专利内术语表

| 术语 | 本文含义/位置 | 备注 |
| --- | --- | --- |
| 超像素 | 说明书 0006 定义为按灰度相似性与空间邻近性划分出的图像块 | 来自说明书定义句 |
| 邻接图 | 说明书 0007 定义为以超像素节点为顶点、空间相邻关系为边的无向图 | 来自说明书定义句 |

## 六、特征—说明书—附图对照

| 特征 | 说明书位置 | 附图 |
| --- | --- | --- |
| F1 | 说明书 0011 | 图1 |
| F2 | 说明书 0012 | 图1 |
| F3 | 说明书 0013 | 图2 |
| F4 | 说明书 0014 | 图2 |
"""

NARRATIVE_MD = """## 一、一句话

把超声图像先切成超像素再连成图，用图卷积判断每块是不是病灶，从而在保持精度的同时压掉逐像素分割的冗余计算。

## 二、连贯叙事

问题：[0002] 指出现有逐像素分割网络忽略相邻组织区域之间的上下文关系。
思路：改以超像素为处理单元，把空间相邻关系显式建成一张图。
怎么做：按说明书 0011 至说明书 0014 的四步依次划分、建图、聚合与归并。
效果：说明书 0003 记载的假阳性与计算冗余问题因此得到缓解。

## 七、和现有技术的差别

说明书 0002 描述的现有技术逐像素独立分类，本方案改为区域级建模并显式表达区域间关系。

## 八、阅读建议

1. 先读说明书 0011-0014 的四个步骤，把主流程走一遍。
2. 关注权2 与权3 对划分算法与节点属性的细化方向。
3. 若考虑规避，可从权1 的邻接图构建环节入手。本方案十分先进。

## 九、技术应用场景

> [!grounding] 应用场景
>
> | 场景/模块 | 大白话 | 专利内依据 |
> | --- | --- | --- |
> | 超声诊断工作站 | 医生在设备上一键得到病灶轮廓 | 说明书 0011 |
"""

APPENDIX_MD = """## 十、附录：行业坐标与公开线索

### A. IPC 行业坐标

本案主分类号落在 A61B 与 G06T，对应生物医药与医疗器械行业的影像诊断环节。来源：离线 IPC 行业词表。

### B. 公开检索线索

> [!warning]- 公开检索线索
>
> - 线索：超声诊断工作站 — 置信度：中 — 理由：说明书 0011 描述的预处理环节与工作站流程吻合；推断自公开文本，未联网核验。
> - 线索：说明书 0099 提到的部署形态（该编号为脚本化测试用的越界引用）；推断自公开文本，未联网核验。

## 十一、免责声明

本解读仅供技术理解辅助，不构成法律意见；专利保护范围以官方法律文本为准。重大决策请咨询专利代理师/律师。
"""

# 自检第一轮：删掉评价性语言（锚点 + 逐字片段，服务端确定性应用）
SELF_CHECK_1: dict[str, Any] = {
    "verdict": "fail",
    "patches": [
        {
            "checklist_item": "标题纪律：各节与卡片标题只用简洁名称",
            "section": "八",
            "anchor_before": "可从权1 的邻接图构建环节入手。",
            "original_excerpt": "本方案十分先进。",
            "replacement": "",
            "reason": "解读报告不应出现评价性语言。",
        }
    ],
    "unresolved": [],
}

SELF_CHECK_2: dict[str, Any] = {
    "verdict": "pass",
    "patches": [],
    "unresolved": [
        {
            "checklist_item": "公开线索",
            "issue": "第十节 B 的第二条线索缺少可核验的说明书依据",
            "needs": "missing_source",
        }
    ],
}

STREAM_SCRIPT: dict[str, str] = {
    "structure": STRUCTURE_MD,
    "narrative": NARRATIVE_MD,
    "appendix": APPENDIX_MD,
}

STRUCT_SCRIPT: dict[str, Any] = {
    "tree": CLAIM_TREE_BAD,
    "tree.repair1": CLAIM_TREE_FIXED,
    "deltas": CLAIM_DELTAS_BAD,
    "deltas.repair1": CLAIM_DELTAS_FIXED,
    "grounding": GROUNDING,
    "self_check1": SELF_CHECK_1,
    "self_check2": SELF_CHECK_2,
}


class FakeLLM:
    """按 step_key 特征回放预置产物；未脚本化的调用直接抛错，便于暴露流程变化。"""

    def __init__(self) -> None:
        self.calls: list[str] = []

    @staticmethod
    def _tag(kwargs: dict[str, Any]) -> str:
        step_key = str(kwargs.get("step_key") or "")
        return step_key.split(".", 1)[1] if "." in step_key else step_key

    async def chat(self, messages, **kwargs) -> str:
        tag = self._tag(kwargs)
        self.calls.append(f"chat:{tag}")
        if tag in STREAM_SCRIPT:
            return STREAM_SCRIPT[tag]
        raise AssertionError(f"FakeLLM 未脚本化的 chat 调用：{tag}")

    async def chat_stream(self, messages, **kwargs):
        tag = self._tag(kwargs)
        self.calls.append(f"stream:{tag}")
        text = STREAM_SCRIPT.get(tag)
        if text is None:
            raise AssertionError(f"FakeLLM 未脚本化的流式调用：{tag}")
        for i in range(0, len(text), 64):
            yield text[i : i + 64]

    async def structured(self, messages, model_cls, **kwargs):
        tag = self._tag(kwargs)
        self.calls.append(f"struct:{tag}")
        data = STRUCT_SCRIPT.get(tag)
        if data is None:
            raise AssertionError(f"FakeLLM 未脚本化的结构化调用：{tag}（模型 {model_cls.__name__}）")
        return model_cls.model_validate(data)


@pytest.fixture()
def fake_llm(monkeypatch: pytest.MonkeyPatch) -> FakeLLM:
    from app.services import llm

    fake = FakeLLM()
    monkeypatch.setattr(llm, "chat", fake.chat)
    monkeypatch.setattr(llm, "chat_stream", fake.chat_stream)
    monkeypatch.setattr(llm, "structured", fake.structured)
    return fake


@pytest.fixture()
def fake_fetch(monkeypatch: pytest.MonkeyPatch):
    """把取证换成「返回本地样例 PDF」；`ok=False` 时模拟全部取证源失败。"""
    from app.services import patent_fetch

    def install(ok: bool = True):
        async def _fetch(pub_no: str, **_kwargs) -> patent_fetch.FetchResult:
            pub = patent_fetch.normalize_pub_no(pub_no)
            if ok:
                return patent_fetch.FetchResult(
                    pub_no=pub,
                    ok=True,
                    pdf=PDF_BYTES,
                    source_id="google_patentimages_cdn",
                    url=f"https://patentimages.storage.googleapis.com/ab/cd/{pub}.pdf",
                    attempts=[
                        patent_fetch.FetchAttempt("google_patents_page", "", True, ""),
                        patent_fetch.FetchAttempt("google_patentimages_cdn", "", True, ""),
                    ],
                )
            return patent_fetch.FetchResult(
                pub_no=pub,
                ok=False,
                error="google_patents_page：HTTP 403；cnipa_epub：详情页未提供可脚本化的全文 PDF",
                attempts=[
                    patent_fetch.FetchAttempt("google_patents_page", "", False, "HTTP 403"),
                ],
            )

        monkeypatch.setattr(patent_fetch, "fetch_patent_pdf", _fetch)

    return install


# ---------------------------------------------------------------------------
# 三、夹具与驱动工具
# ---------------------------------------------------------------------------



def _new_case(client: TestClient, title: str) -> str:
    resp = client.post("/api/v1/cases", json={"module": "reader", "title": title})
    assert resp.status_code == 201, resp.text
    return resp.json()["id"]


def _upload_pdf(client: TestClient, case_id: str) -> str:
    resp = client.post(
        f"/api/v1/cases/{case_id}/files",
        files={"files": (f"{PUB_NO}.pdf", PDF_BYTES, "application/pdf")},
    )
    assert resp.status_code == 201, resp.text
    return resp.json()[0]["file"]["id"]


def _steps(case_id: str):
    from app.db import database as db
    from app.pipelines import registry

    row = db.query_one("SELECT * FROM cases WHERE id=?", (case_id,))
    return registry.build_steps(dict(row))


def _case_row(case_id: str) -> dict[str, Any]:
    from app.db import database as db

    return dict(db.query_one("SELECT * FROM cases WHERE id=?", (case_id,)))


def _case_state(case_id: str) -> dict[str, Any]:
    return json.loads(_case_row(case_id)["state_json"] or "{}")


def _failures(case_id: str) -> str:
    from app.db import database as db

    rows = db.query_all(
        "SELECT step_key, status, error FROM pipeline_runs WHERE case_id=? AND status='failed'",
        (case_id,),
    )
    return "；".join(f"{r['step_key']}: {r['error']}" for r in rows) or "无失败步骤记录"


def _artifacts(case_id: str, kind: str) -> list[dict[str, Any]]:
    from app.db import database as db

    rows = db.query_all(
        "SELECT * FROM artifacts WHERE case_id=? AND kind=? ORDER BY version ASC", (case_id, kind)
    )
    return [dict(r) for r in rows]


async def _drive(case_id: str, task: asyncio.Task, answers: dict[str, Any]) -> list[str]:
    """轮询门控并按 step_key 提交预置答复，直到流水线结束；返回遇到的门控顺序。"""
    from app.pipelines import engine

    seen: list[str] = []
    while not task.done():
        pending = engine.get_pending(case_id)
        if pending is not None:
            step_key = pending["step_key"]
            if step_key not in seen:
                seen.append(step_key)
            try:
                engine.submit_input(case_id, step_key, answers.get(step_key, {}))
            except engine.PipelineConflictError:
                pass
        await asyncio.sleep(0.05)
    await task
    return seen


# 跨用例共享主流程结果
_run: dict[str, Any] = {}


# ---------------------------------------------------------------------------
# 四、主流程：五步跑通
# ---------------------------------------------------------------------------


async def test_reader_pipeline_runs_to_deliver(
    client: TestClient, fake_llm: FakeLLM, fake_fetch
) -> None:
    """公开号 → 取证 → 五步跑通，案件 completed。"""
    from app.pipelines import engine

    fake_fetch(ok=True)
    case_id = _new_case(client, "专利解读端到端测试")
    task = engine.start(
        case_id, _steps(case_id), run_group="initial", start_payload={"pub_no": PUB_NO}
    )
    gates = await asyncio.wait_for(_drive(case_id, task, {"claim_tree": {"skip": True}}), timeout=600)

    assert _case_row(case_id)["status"] == "completed", _failures(case_id)
    assert gates == ["claim_tree"], "权要树消歧门控应当且只当被触发一次"
    _run["case_id"] = case_id
    _run["state"] = _case_state(case_id)
    _run["calls"] = list(fake_llm.calls)


def test_five_steps_all_done() -> None:
    """五个步骤都留下了 done 的 pipeline_runs 行。"""
    from app.db import database as db

    rows = db.query_all(
        "SELECT step_key, status FROM pipeline_runs WHERE case_id=? ORDER BY started_at",
        (_run["case_id"],),
    )
    done = {r["step_key"] for r in rows if r["status"] == "done"}
    assert done == {"acquire", "claim_tree", "claim_deltas", "note", "lint_deliver"}


def test_sse_event_contract() -> None:
    """SSE 契约事件齐全：step_status / llm_done(doc) / interaction_required / 交付物 / pipeline_done。"""
    from app.db import database as db

    rows = db.query_all(
        "SELECT step_key, content, meta_json FROM messages WHERE case_id=? ORDER BY seq",
        (_run["case_id"],),
    )
    named = [(json.loads(r["meta_json"])["event"], json.loads(r["content"])) for r in rows]
    events = {name for name, _ in named}
    assert {
        "step_status",
        "log",
        "case_title",
        "interaction_required",
        "llm_done",
        "artifact_created",
        "doc_version",
        "pipeline_done",
    } <= events

    dones = [data for name, data in named if name == "llm_done"]
    assert any(d.get("channel") == "doc" and d.get("doc_id") == "reader_note" for d in dones)
    assert any(d.get("channel") == "chat" for d in dones)

    interactions = [data for name, data in named if name == "interaction_required"]
    assert len(interactions) == 1
    assert interactions[0]["step_key"] == "claim_tree" and interactions[0]["kind"] == "claim_tree"
    assert interactions[0]["schema"]["properties"]["claims"]["type"] == "array"
    assert [c["no"] for c in interactions[0]["default"]["claims"]] == [1, 2, 3, 4, 5]

    finals = [data for name, data in named if name == "pipeline_done"]
    assert finals[-1] == {"run_group": "initial", "status": "done"}


def test_acquire_structure_and_type() -> None:
    """acquire：PDF→md 复用 convert，切分出权项/段落号/附图，种类码判为发明。"""
    state = _run["state"]
    meta = state["reader_meta"]
    assert meta["pub_no"] == PUB_NO
    assert meta["title"] == PATENT_TITLE
    assert meta["type"] == "invention" and meta["type_label"] == "发明"
    assert state["reader_type"]["source"] == "pub_kind_code"   # 种类码主路径，未走 LLM 兜底
    assert "struct:type_hook" not in _run["calls"]
    assert meta["ipc"][:1] == ["G06T 7/11"]

    summary = state["reader_structure"]
    assert summary["claim_count"] == 5
    assert summary["paragraph_count"] == 15
    assert summary["figure_count"] == 2
    assert summary["evidence_scope"] == "full_text"

    # 结构切分产物落盘（全文不进 state_json）
    path = Path(state["reader_structure_path"])
    assert path.is_file()
    data = json.loads(path.read_text(encoding="utf-8"))
    assert data["paragraphs"]["0006"].startswith("所述超像素划分是指")
    assert "full_text" not in state["reader_structure"]        # state 里只有轻量摘要

    # 案件标题/类型已回填
    row = _case_row(_run["case_id"])
    assert row["title"] == PATENT_TITLE
    assert row["patent_type"] == "invention"


# ---------------------------------------------------------------------------
# 五、权要树 lint 与 REPAIR
# ---------------------------------------------------------------------------


def test_claim_tree_lint_catches_illegal_refs() -> None:
    """树 lint 捕获「引用号大于自身权号」「引用不存在的权项」「多引再被引」「环」。"""
    from app.models.reader import ClaimTree
    from app.pipelines.reader import lint_claim_tree

    bad = ClaimTree.model_validate(CLAIM_TREE_BAD)
    findings = lint_claim_tree(bad)
    errors = [f for f in findings if f["severity"] == "error"]
    assert any(f["rule"] == "ref_order" and f["claim"] == 3 for f in errors), findings

    # 引用不存在的权项（编号小于自身，但树里根本没有这一项）
    ghost = ClaimTree.model_validate(
        {"claims": [_CLAIM_NODES[0], {**_CLAIM_NODES[4], "type": "dependent", "refs": [4]}]}
    )
    assert any(f["rule"] == "ref_missing" for f in lint_claim_tree(ghost))

    # 多项引用的权利要求被再次引用（权4 多引，权5 又引权4）
    chained = ClaimTree.model_validate(
        {
            "claims": [
                *_CLAIM_NODES[:4],
                {
                    **_CLAIM_NODES[4],
                    "type": "dependent",
                    "refs": [4],
                    "ref_mode": "single",
                },
            ]
        }
    )
    assert any(f["rule"] == "multi_ref_chain" for f in lint_claim_tree(chained))

    # 循环引用（用户在门控里编辑出来的坏树）
    cyclic = ClaimTree.model_validate(
        {
            "claims": [
                {**_CLAIM_NODES[0], "type": "dependent", "refs": [2], "ref_mode": "single"},
                {**_CLAIM_NODES[1], "refs": [1]},
            ]
        }
    )
    assert any(f["rule"] == "cycle" for f in lint_claim_tree(cyclic))

    # 合法树无 error
    assert not [f for f in lint_claim_tree(ClaimTree.model_validate(CLAIM_TREE_FIXED))
                if f["severity"] == "error"]


def test_claim_tree_repair_and_gate() -> None:
    """首版非法引用触发一次 REPAIR；ambiguities 非空触发门控，跳过时按 alternative 处理。"""
    state = _run["state"]
    assert "struct:tree.repair1" in _run["calls"]
    assert state["claim_tree_repairs"] == 1
    assert state["claim_tree_resolved"] is False               # 用户跳过了消歧

    tree = state["claim_tree"]
    by_no = {c["no"]: c for c in tree["claims"]}
    assert by_no[3]["refs"] == [1]                             # REPAIR 修正了越界引用
    assert by_no[4]["ref_mode"] == "alternative"               # 跳过 → 默认按择一引用
    assert not [f for f in state["claim_tree_lint"] if f["severity"] == "error"]


def test_merge_user_tree_applies_edits() -> None:
    """用户在门控里改了引用方式时，编辑结果被合并回权要树。"""
    from app.models.reader import ClaimTree
    from app.pipelines.reader import _merge_user_tree

    tree = ClaimTree.model_validate(CLAIM_TREE_FIXED)
    merged, resolved = _merge_user_tree(
        tree, {"claims": [{"no": 4, "refs": [2], "ref_mode": "single"}]}
    )
    assert resolved is True
    assert merged.by_no()[4].refs == [2]
    assert merged.by_no()[4].ref_mode == "single"
    assert merged.by_no()[1].type == "independent"             # 未提交的权项保持原样


# ---------------------------------------------------------------------------
# 六、白话增量长度校验与 REPAIR
# ---------------------------------------------------------------------------


def test_claim_deltas_length_repair() -> None:
    """首版超 40 字触发一次 REPAIR，终版全部落在 12–40 字。"""
    from app.models.reader import DELTA_MAX_CHARS, DELTA_MIN_CHARS, ClaimDeltas, ClaimTree
    from app.pipelines.reader import lint_claim_deltas

    state = _run["state"]
    assert "struct:deltas.repair1" in _run["calls"]
    assert state["claim_deltas_repairs"] == 1

    deltas = ClaimDeltas.model_validate(state["claim_deltas"])
    assert {d.claim for d in deltas.deltas} == {1, 2, 3, 4, 5}
    for item in deltas.deltas:
        assert DELTA_MIN_CHARS <= item.length <= DELTA_MAX_CHARS, item.delta
    assert not [f for f in state["claim_deltas_lint"] if f["severity"] == "error"]

    # 长度校验本身：超长 / 过短 / 套话都要被抓住
    tree = ClaimTree.model_validate(CLAIM_TREE_FIXED)
    findings = lint_claim_deltas(tree, ClaimDeltas.model_validate(CLAIM_DELTAS_BAD))
    assert any(f["rule"] == "too_long" and f["claim"] == 4 for f in findings)

    short = ClaimDeltas.model_validate({"deltas": [{"claim": 1, "delta": "太短了"}]})
    rules = {f["rule"] for f in lint_claim_deltas(tree, short)}
    assert "too_short" in rules and "missing" in rules

    cliche = ClaimDeltas.model_validate(
        {"deltas": [{"claim": 1, "delta": "根据权利要求1所述的方法，其特征在于增加了一步"}]}
    )
    assert any(f["rule"] == "cliche" for f in lint_claim_deltas(tree, cliche))


# ---------------------------------------------------------------------------
# 七、报告：11 节齐全 + 引用 lint + 自检补丁
# ---------------------------------------------------------------------------


def test_report_has_eleven_sections(client: TestClient) -> None:
    """report.json：meta + 11 节，节标题与模板逐字一致，第三节含平台装配的树形表。"""
    resp = client.get(f"/api/v1/cases/{_run['case_id']}/reader/report")
    assert resp.status_code == 200, resp.text
    payload = resp.json()
    report = payload["report"]

    assert report["meta"]["pub_no"] == PUB_NO
    assert report["meta"]["type"] == "invention"
    assert report["meta"]["ipc"]

    sections = report["sections"]
    assert len(sections) == 11
    assert [s["id"] for s in sections] == [f"s{i}" for i in range(1, 12)]
    assert [s["title"] for s in sections] == [
        "一、一句话",
        "二、连贯叙事",
        "三、权利要求树",
        "四、独立权利要求精读",
        "五、专利内术语表",
        "六、特征—说明书—附图对照",
        "七、和现有技术的差别",
        "八、阅读建议",
        "九、技术应用场景",
        "十、附录：行业坐标与公开线索",
        "十一、免责声明",
    ]
    assert all(s["blocks"] for s in sections), "每一节都应有可渲染的 block"

    by_id = {s["id"]: s for s in sections}
    # 著录项卡片
    assert by_id["s1"]["blocks"][0]["type"] == "callout-meta"
    assert PUB_NO in by_id["s1"]["blocks"][0]["content"]["markdown"]
    # 第三节：单一主展示 = 平台装配的 claim_tree block（不再有裸表格）
    types_s3 = [b["type"] for b in by_id["s3"]["blocks"]]
    assert "claim_tree" in types_s3 and "table" not in types_s3
    tree_block = next(b for b in by_id["s3"]["blocks"] if b["type"] == "claim_tree")
    assert tree_block["content"]["summary"].startswith("共 5 项 · 独立 2 / 从属 3")
    rows = tree_block["content"]["rows"]
    assert [r["no"] for r in rows] == [1, 2, 3, 4, 5]
    assert rows[0]["structure"] == "◆" and rows[0]["delta"]
    assert rows[3]["delta"] == "限定图卷积网络的层数固定为三层"
    # 第四节：独立权卡片 + 特征表
    types_s4 = [b["type"] for b in by_id["s4"]["blocks"]]
    assert types_s4.count("callout-claim") == 2 and "table" in types_s4
    # 第九节：grounding 卡片；第十节：warning 卡片
    assert any(b["type"] == "callout-grounding" for b in by_id["s9"]["blocks"])
    assert any(b["type"] == "callout-warning" for b in by_id["s10"]["blocks"])
    # 第十一节：免责声明逐字
    assert "不构成法律意见" in json.dumps(by_id["s11"]["blocks"], ensure_ascii=False)


def test_paragraph_reference_lint(client: TestClient) -> None:
    """说明书引用 lint：裸 [0002] 改写、区间归一 en-dash、越界段落号就地标注。"""
    resp = client.get(f"/api/v1/cases/{_run['case_id']}/reader/report")
    markdown = resp.json()["markdown"]

    assert "[0002]" not in markdown                       # 禁裸段落号
    assert "说明书 0002" in markdown                       # 已改写为规范格式
    assert "说明书 0011–0014" in markdown                  # 半角连字符归一为 en-dash
    assert "说明书 0011-0014" not in markdown
    assert "说明书 0099（该段落号未在本案说明书中检索到）" in markdown
    assert markdown.count("（该段落号未在本案说明书中检索到）") == 1

    rules = {f["rule"] for f in resp.json()["lint"]}
    assert {"bare_paragraph", "paragraph_missing", "range_dash"} <= rules

    # 落地线索里锚定到不存在段落号的那条被服务端丢弃
    grounding = _run["state"]["grounding"]
    assert [c["scene"] for c in grounding["clues"]] == ["超声诊断工作站"]
    assert any(f["rule"] == "grounding_anchor" for f in _run["state"]["grounding_lint"])


def test_self_check_patch_applied(client: TestClient) -> None:
    """AUDIT 补丁经 services/patches 确定性应用；未决事项不入正文。"""
    state = _run["state"]
    markdown = state["report_markdown"]
    assert "本方案十分先进。" not in markdown              # 补丁删掉了评价性语言
    assert "可从权1 的邻接图构建环节入手。" in markdown      # 只删该句，锚点原文保留
    assert state["self_check"]["applied"] == 1
    assert state["self_check"]["rounds"] == 2
    assert any("缺少可核验的说明书依据" in u for u in state["self_check"]["unresolved"])

    resp = client.get(f"/api/v1/cases/{_run['case_id']}/reader/report")
    assert any("缺少可核验的说明书依据" in u for u in resp.json()["unresolved"])


# ---------------------------------------------------------------------------
# 八、交付物：reader_note_md 落盘 + docx 导出
# ---------------------------------------------------------------------------


def test_reader_note_artifact_and_docx_export(client: TestClient) -> None:
    """reader_note_md 落盘（版本化命名），并可经 /artifacts/{id}/export 出 docx。"""
    case_id = _run["case_id"]
    notes = _artifacts(case_id, "reader_note_md")
    assert notes, "未产出 reader_note_md 交付物"
    latest = notes[-1]
    assert latest["filename"].startswith(f"{PATENT_TITLE}_")
    text = disk_path(latest["stored_path"]).read_text(encoding="utf-8")
    assert text.startswith(f"# 专利解读：{PATENT_TITLE}")
    assert text.count("\n## ") == 11                       # 11 节齐全
    assert text == _run["state"]["report_markdown"]

    exported = client.post(f"/api/v1/artifacts/{latest['id']}/export", json={"format": "docx"})
    assert exported.status_code == 200, exported.text
    art = exported.json()
    assert art["kind"] == "reader_note_docx"
    assert art["source_artifact_id"] == latest["id"]

    from docx import Document

    doc = Document(disk_path(art["stored_path"]))
    texts = [p.text for p in doc.paragraphs]
    assert any("专利解读" in t for t in texts)
    assert any("免责声明" in t for t in texts)


def test_claim_tree_api(client: TestClient) -> None:
    """GET /reader/claim-tree：树 + 白话增量 + lint。"""
    resp = client.get(f"/api/v1/cases/{_run['case_id']}/reader/claim-tree")
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert [c["no"] for c in data["tree"]["claims"]] == [1, 2, 3, 4, 5]
    assert data["summary"].startswith("共 5 项 · 独立 2 / 从属 3")
    assert {d["claim"] for d in data["deltas"]} == {1, 2, 3, 4, 5}
    assert data["ambiguities_resolved"] is False
    assert isinstance(data["lint"], list)


def test_reader_api_404_on_other_case(client: TestClient) -> None:
    """未跑过解读的案件 → 404；非解读案件 → 400。"""
    empty = _new_case(client, "空解读案件")
    assert client.get(f"/api/v1/cases/{empty}/reader/report").status_code == 404
    assert client.get(f"/api/v1/cases/{empty}/reader/claim-tree").status_code == 404

    other = client.post("/api/v1/cases", json={"module": "paper2patent", "title": "非解读案件"})
    other_id = other.json()["id"]
    assert client.get(f"/api/v1/cases/{other_id}/reader/report").status_code == 400


# ---------------------------------------------------------------------------
# 九、取证失败 → 手工上传门控
# ---------------------------------------------------------------------------


async def test_fetch_failure_falls_back_to_upload_gate(
    client: TestClient, fake_llm: FakeLLM, fake_fetch
) -> None:
    """全部取证源失败时不抛错，而是门控请用户上传 PDF（附可点链接）。"""
    from app.pipelines import engine

    fake_fetch(ok=False)
    case_id = _new_case(client, "取证失败兜底测试")
    file_id = _upload_pdf(client, case_id)

    prompts: dict[str, Any] = {}

    async def drive() -> list[str]:
        seen: list[str] = []
        while not task.done():
            pending = engine.get_pending(case_id)
            if pending is not None:
                step_key = pending["step_key"]
                if step_key not in seen:
                    seen.append(step_key)
                    prompts[step_key] = pending
                answer = {"acquire": {"file_id": file_id}, "claim_tree": {"skip": True}}
                try:
                    engine.submit_input(case_id, step_key, answer.get(step_key, {}))
                except engine.PipelineConflictError:
                    pass
            await asyncio.sleep(0.05)
        await task
        return seen

    task = engine.start(
        case_id, _steps(case_id), run_group="initial", start_payload={"pub_no": PUB_NO}
    )
    gates = await asyncio.wait_for(drive(), timeout=600)

    assert _case_row(case_id)["status"] == "completed", _failures(case_id)
    assert gates == ["acquire", "claim_tree"]

    gate = prompts["acquire"]
    assert gate["kind"] == "form"
    assert "HTTP 403" in gate["prompt"]                     # 如实交代失败原因
    links = gate["default"]["links"]
    assert any("patents.google.com" in item["url"] for item in links)
    assert any("cnipa.gov.cn" in item["url"] for item in links)

    state = _case_state(case_id)
    assert state["reader_source"]["manual_upload"] is True
    assert state["reader_meta"]["title"] == PATENT_TITLE
    assert _artifacts(case_id, "reader_note_md")


# ---------------------------------------------------------------------------
# 十、单元测试：取证 / 结构切分 / 装配
# ---------------------------------------------------------------------------


async def test_fetch_patent_pdf_never_raises(monkeypatch: pytest.MonkeyPatch) -> None:
    """网络全挂时 fetch_patent_pdf 返回 ok=False（不抛异常），并逐源记录失败原因。"""
    import httpx

    from app.services import patent_fetch

    class DeadClient:
        async def get(self, url, **_kwargs):
            raise httpx.ConnectError("网络不可达")

        async def aclose(self) -> None:
            return None

    result = await patent_fetch.fetch_patent_pdf("CN999999999A", client=DeadClient())
    assert result.ok is False and result.pdf is None
    assert result.error
    assert {a.source_id for a in result.attempts} >= {"google_patents_page", "cnipa_epub"}
    assert all(a.ok is False for a in result.attempts)
    # to_dict 可直接入 state_json（不含 PDF 字节）
    assert result.to_dict()["size"] == 0


def test_fetch_sources_priority_and_links() -> None:
    """取证源按 patent_pdf_sources.yaml 的 priority 排序；手工链接可点。"""
    from app.services import patent_fetch

    sources = patent_fetch.load_sources()
    ids = [s["id"] for s in sorted(sources["sources"], key=lambda s: s["priority"])]
    assert ids.index("google_patents_page") < ids.index("cnipa_epub")

    assert patent_fetch.normalize_pub_no(" cn 209861402 u ") == "CN209861402U"
    assert patent_fetch.known_cdn_url("CN209861402U").endswith("CN209861402U.pdf")
    links = patent_fetch.manual_links("CN209861402U")
    assert links[0]["url"].startswith("https://patents.google.com/patent/")

    html = '<meta name="citation_pdf_url" content="https://patentimages.storage.googleapis.com/x/y/CN1A.pdf">'
    assert patent_fetch.extract_pdf_url(html).endswith("CN1A.pdf")
    assert patent_fetch.extract_pdf_url("<html>没有直链</html>") is None
    assert patent_fetch.looks_like_pdf(b"%PDF-1.7 ...") and not patent_fetch.looks_like_pdf(b"<html>")


def test_structure_parsing_of_sample_pdf() -> None:
    """结构切分：扉页著录项 / 权项 / 段落号索引表 / 附图清单。"""
    import pymupdf

    from app.services import patent_fetch

    with pymupdf.open(stream=PDF_BYTES, filetype="pdf") as doc:
        md = "\n".join(f"## 第 {i} 页\n\n{p.get_text('text')}" for i, p in enumerate(doc, 1))

    st = patent_fetch.parse_patent_md(md)
    assert st.pub_no == PUB_NO
    assert st.title == PATENT_TITLE
    assert st.ipc == ["G06T 7/11", "A61B 8/08"]
    assert st.applicants and "引途医疗" in st.applicants[0]
    assert st.app_no == "202410123456.7" and st.pub_date == "2025-03-14"
    assert "医学图像处理" in st.abstract

    assert [c["no"] for c in st.claims] == [1, 2, 3, 4, 5]
    assert st.claims[0]["text"].startswith("一种超声图像病灶区域分割方法")
    assert "或3所述" in st.claims[3]["text"]

    assert sorted(st.paragraphs) == [f"{i:04d}" for i in range(1, 16)]
    assert st.paragraphs["0002"].startswith("现有超声图像分割方法")
    assert "[0002]" not in st.paragraphs["0002"]

    assert [f["no"] for f in st.figures] == [1, 2]
    assert "流程图" in st.figures[0]["caption"]
    assert st.evidence_scope() == "full_text"

    # 段落号索引与判型链路
    from app.tools.patent_type import infer_patent_type_from_pub

    assert infer_patent_type_from_pub(st.pub_no) == "invention"


def test_ipc_hints_matching() -> None:
    """IPC 行业提示：长前缀优先 → 关键词 → DEFAULT 兜底。"""
    from app.pipelines.reader import match_ipc_hints

    picked = match_ipc_hints(["H04L 12/00", "H04W 4/00"], "")
    assert picked[0]["ipc_prefix"] == "H04L"                 # H04L 比 H04 更长，优先

    assert match_ipc_hints(["A61B 8/08"], "")[0]["industry"] == "生物医药与医疗器械"
    assert match_ipc_hints([], "本方案用于锂离子电池隔膜")[0]["ipc_prefix"] == "H01M"
    assert match_ipc_hints([], "毫无关联的文本")[0]["ipc_prefix"] == "DEFAULT"


def test_markdown_blocks_and_claim_tree_table() -> None:
    """markdown → block 序列；权要树表由平台按 claim_tree + claim_deltas 装配。"""
    from app.models.reader import ClaimDeltas, ClaimTree
    from app.pipelines.reader import build_claim_tree_block, markdown_to_blocks

    blocks = markdown_to_blocks(
        "正文一段。\n\n"
        "> [!grounding] 应用场景\n> 卡片正文。\n\n"
        "| A | B |\n| --- | --- |\n| 1 | 2 |\n\n"
        "```mermaid\nflowchart TD\n  A --> B\n```\n\n"
        "> [!figure] 图注\n> 未映射的 callout 退回 markdown。\n"
    )
    types = [b.type for b in blocks]
    assert types == ["markdown", "callout-grounding", "table", "mermaid", "markdown"]
    assert blocks[1].content == {"title": "应用场景", "markdown": "卡片正文。"}
    assert blocks[2].content["headers"] == ["A", "B"]
    assert blocks[2].content["rows"] == [["1", "2"]]
    assert blocks[3].content.startswith("flowchart TD")

    tree = ClaimTree.model_validate(CLAIM_TREE_FIXED)
    block = build_claim_tree_block(tree, ClaimDeltas.model_validate(CLAIM_DELTAS_FIXED))
    assert block.type == "claim_tree"
    glyphs = {r["no"]: r["structure"] for r in block.content["rows"]}
    assert glyphs[1] == "◆" and glyphs[5] == "◆"             # 两项独立权
    assert glyphs[3].endswith("└─")                          # 权1 下最后一个直接从属权
    assert glyphs[4].startswith("│")                         # 二级从属权有缩进
    assert "| 结构 | 权 | 本项新增 |" in block.content["markdown"]


async def test_mermaid_gate_degrades_unrenderable(monkeypatch: pytest.MonkeyPatch) -> None:
    """mermaid 门禁：能渲染的原样保留，渲染不过的降级为编号文字清单并记 warning。"""
    from app.pipelines.reader import lint_mermaid
    from app.services import disclosure_build

    md = "## 三、权利要求树\n\n```mermaid\nflowchart TD\n  C1[\"权1\"] --> C2[\"权2\"]\n```\n"

    async def ok(_code: str):
        return b"\x89PNG", None

    monkeypatch.setattr(disclosure_build, "render_mermaid", ok)
    kept, degraded, findings = await lint_mermaid(md)
    assert kept == md and degraded == 0 and findings == []

    async def broken(_code: str):
        return None, "浏览器不可用"

    monkeypatch.setattr(disclosure_build, "render_mermaid", broken)
    text, degraded, findings = await lint_mermaid(md)
    assert degraded == 1
    assert "```mermaid" not in text and "图待补" in text
    assert "权1" in text and "权2" in text                    # 降级清单保留节点标签
    assert findings[0]["rule"] == "mermaid"

    # 没有 mermaid 围栏时是纯直通（不触碰浏览器）
    assert await lint_mermaid("## 三、权利要求树\n\n共 2 项。\n") == (
        "## 三、权利要求树\n\n共 2 项。\n",
        0,
        [],
    )


def test_writer_drawn_tree_table_is_not_duplicated() -> None:
    """撰写调用若自己也画了树形表，装配时剔除，保证「不得双份主展示」。"""
    from app.models.reader import ClaimDeltas, ClaimTree, ReaderMeta
    from app.pipelines.reader import assemble_markdown, build_report, strip_tree_tables

    rogue = (
        "共 5 项 · 独立 2 / 从属 3。\n\n"
        "| 结构 | 权 | 本项新增 |\n| --- | ---: | --- |\n| `◆` | 1 | 模型自己画的 |\n\n"
        "| 别的表 | 值 |\n| --- | --- |\n| A | 1 |\n"
    )
    stripped = strip_tree_tables(rogue)
    assert "模型自己画的" not in stripped and "别的表" in stripped

    tree = ClaimTree.model_validate(CLAIM_TREE_FIXED)
    deltas = ClaimDeltas.model_validate(CLAIM_DELTAS_FIXED)
    sections = {sid: "占位正文。" for sid in [f"s{i}" for i in range(1, 12)]}
    sections["s3"] = rogue
    markdown = assemble_markdown(
        ReaderMeta(pub_no=PUB_NO, title=PATENT_TITLE), sections, tree, deltas, run_group="initial"
    )
    assert markdown.count("| 结构 | 权 | 本项新增 |") == 1
    assert "模型自己画的" not in markdown

    report = build_report(markdown, ReaderMeta(pub_no=PUB_NO), tree, deltas)
    s3 = report.section("s3")
    assert [b.type for b in s3.blocks].count("claim_tree") == 1
    assert any(b.type == "table" for b in s3.blocks)          # 非树形表的表格照常保留


def test_split_writer_sections() -> None:
    """撰写调用输出按 `## 序号、名称` 切节；十一 不被误切成 十。"""
    from app.pipelines.reader import split_writer_sections

    sections = split_writer_sections(APPENDIX_MD)
    assert set(sections) == {"s10", "s11"}
    assert sections["s10"].startswith("### A. IPC 行业坐标")
    assert sections["s11"].startswith("本解读仅供技术理解辅助")
    assert set(split_writer_sections(STRUCTURE_MD)) == {"s3", "s4", "s5", "s6"}
    assert set(split_writer_sections(NARRATIVE_MD)) == {"s1", "s2", "s7", "s8", "s9"}
