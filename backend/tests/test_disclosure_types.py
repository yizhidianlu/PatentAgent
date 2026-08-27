"""模块 A（专利交底书 · 实用新型 / 外观设计分支）端到端测试。

与 `test_disclosure.py` 同一套路数，不依赖真实 LLM 与真实网络：

- `FakeLLM` 按 `step_key` 后缀脚本化回放预置产物（chat_stream / chat / structured 三入口）；
- `cnipa.search` 被替换为「检索失败」的测试桩 → 走查新失败门控的「跳过」分支
  （1.1 如实写明未检索，平台不编造检索结果）；
- 其余环节**全部真跑**：A3b 事实合同填表与 figure_plan 服务端硬规则、附图入文规则、
  第五章装置书式 lint（正则 + AUDIT 双查）与 REPAIR、章节 lint 束、md→docx 真子进程。

覆盖点：

1. 九步流水线（多一个 `schema_fill`）跑到 deliver；`schema_fill` 门控回填能覆盖 state；
2. figure_plan 入文规则：CAD 永不入文（连用户手动勾选都要被打回）、低分场景实拍不入文、
   实用新型只收线稿、外观实拍 + 线稿都入文；
3. 章节结构分别符合各自模板（实用新型 3.1 总体构成 / 3.2 连接与配合；外观三、视图说明）；
4. 实用新型第五章书式 lint 生效：首版写成方法书式 → `u4.repair1` 改写为装置书式；
5. 发明分支不受影响（步骤表仍是八步，`schema_fill` 对发明直通）。
"""

from __future__ import annotations

import asyncio
import json
import struct
import zlib
from pathlib import Path
from typing import Any

import pytest
from fastapi.testclient import TestClient

# ---------------------------------------------------------------------------
# 通用夹具工具（图片、FakeLLM、检索桩、门控驱动）
# ---------------------------------------------------------------------------

SENSITIVE = ["引途医疗科技有限公司", "华东某汽车电子厂", "37 万元"]

LINEART = "lineart_assembly.png"
PHOTO_CLEAN = "photo_clean.png"
PHOTO_SCENE = "photo_scene.png"
CAD = "part_cad.png"


def _png_bytes(width: int = 64, height: int = 48, gray: int = 240) -> bytes:
    """生成一张合法的纯色 RGB PNG（避免测试依赖 Pillow 造图）。"""

    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    raw = b"".join(b"\x00" + bytes([gray, gray, gray]) * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


class FakeLLM:
    """按 step_key 后缀回放预置产物；未脚本化的调用直接抛错，便于暴露流程变化。"""

    def __init__(
        self,
        *,
        stream: dict[str, str],
        chat: dict[str, str] | None = None,
        struct_: dict[str, Any],
    ) -> None:
        self.calls: list[str] = []
        self.stream_script = dict(stream)
        self.chat_script = dict(chat or {})
        self.struct_script = dict(struct_)

    @staticmethod
    def _tag(kwargs: dict[str, Any]) -> str:
        step_key = str(kwargs.get("step_key") or "")
        return step_key.split(".", 1)[1] if "." in step_key else step_key

    async def chat(self, messages, **kwargs) -> str:
        tag = self._tag(kwargs)
        self.calls.append(f"chat:{tag}")
        if tag in self.chat_script:
            return self.chat_script[tag]
        if tag in self.stream_script:
            return self.stream_script[tag]
        raise AssertionError(f"FakeLLM 未脚本化的 chat 调用：{tag}")

    async def chat_stream(self, messages, **kwargs):
        tag = self._tag(kwargs)
        self.calls.append(f"stream:{tag}")
        text = self.stream_script.get(tag)
        if text is None:
            raise AssertionError(f"FakeLLM 未脚本化的流式调用：{tag}")
        for i in range(0, len(text), 64):
            yield text[i : i + 64]

    async def structured(self, messages, model_cls, **kwargs):
        tag = self._tag(kwargs)
        self.calls.append(f"struct:{tag}")
        data = self.struct_script.get(tag)
        if data is None:
            raise AssertionError(f"FakeLLM 未脚本化的结构化调用：{tag}（模型 {model_cls.__name__}）")
        return model_cls.model_validate(data)


def _install_llm(monkeypatch: pytest.MonkeyPatch, fake: FakeLLM) -> FakeLLM:
    from app.services import llm

    monkeypatch.setattr(llm, "chat", fake.chat)
    monkeypatch.setattr(llm, "chat_stream", fake.chat_stream)
    monkeypatch.setattr(llm, "structured", fake.structured)
    return fake


def _install_failing_search(monkeypatch: pytest.MonkeyPatch) -> None:
    """检索桩：恒定失败（本套用例统一走「跳过查新」分支，不编造检索结果）。"""
    import inspect as _inspect

    from app.models.search import SearchResult
    from app.services import cnipa

    async def fake_search(case_id: str, terms, patent_type: str = "invention", **kwargs):
        on_progress = kwargs.get("on_progress")
        if on_progress is not None:
            result = on_progress("failed", "浏览器不可用（测试桩）")
            if _inspect.isawaitable(result):
                await result
        return SearchResult(
            status="failed",
            hits=[],
            error="浏览器不可用（测试桩）",
            terms=cnipa.normalize_terms(terms),
            patent_type=patent_type,
        )

    monkeypatch.setattr(cnipa, "search", fake_search)



def _new_case(client: TestClient, title: str, patent_type: str) -> str:
    resp = client.post(
        "/api/v1/cases", json={"module": "disclosure", "title": title, "patent_type": patent_type}
    )
    assert resp.status_code == 201, resp.text
    return resp.json()["id"]


def _upload(client: TestClient, case_id: str, name: str, payload: bytes, mime: str) -> str:
    resp = client.post(f"/api/v1/cases/{case_id}/files", files={"files": (name, payload, mime)})
    assert resp.status_code == 201, resp.text
    return resp.json()[0]["file"]["id"]


def _upload_images(client: TestClient, case_id: str) -> None:
    """四张候选图：线稿 / 干净实拍 / 场景实拍 / CAD 投影。"""
    for name in (LINEART, PHOTO_CLEAN, PHOTO_SCENE, CAD):
        _upload(client, case_id, name, _png_bytes(), "image/png")


def _steps(case_id: str):
    from app.db import database as db
    from app.pipelines import registry

    row = db.query_one("SELECT * FROM cases WHERE id=?", (case_id,))
    return registry.build_steps(dict(row))


def _case_row(case_id: str) -> dict[str, Any]:
    from app.db import database as db

    return dict(db.query_one("SELECT * FROM cases WHERE id=?", (case_id,)))


def _case_state(case_id: str) -> dict[str, Any]:
    return json.loads(_case_row(case_id)["state_json"] or "{}")


def _failures(case_id: str) -> str:
    from app.db import database as db

    rows = db.query_all(
        "SELECT step_key, status, error FROM pipeline_runs WHERE case_id=? AND status='failed'",
        (case_id,),
    )
    return "；".join(f"{r['step_key']}: {r['error']}" for r in rows) or "无失败步骤记录"


def _artifacts(case_id: str, kind: str) -> list[dict[str, Any]]:
    from app.db import database as db

    rows = db.query_all(
        "SELECT * FROM artifacts WHERE case_id=? AND kind=? ORDER BY version ASC", (case_id, kind)
    )
    return [dict(r) for r in rows]


async def _drive(case_id: str, task: asyncio.Task, answers: dict[str, Any]) -> list[dict[str, Any]]:
    """轮询门控并提交预置答复（值可为 callable(pending)→payload）；返回门控出现顺序。"""
    from app.pipelines import engine

    seen: list[dict[str, Any]] = []
    while not task.done():
        pending = engine.get_pending(case_id)
        if pending is not None:
            step_key = pending["step_key"]
            if not any(s["step_key"] == step_key and s["kind"] == pending["kind"] for s in seen):
                seen.append({"step_key": step_key, "kind": pending["kind"]})
            answer = answers.get(step_key, {})
            if callable(answer):
                answer = answer(pending)
            try:
                engine.submit_input(case_id, step_key, answer)
            except engine.PipelineConflictError:
                pass
        await asyncio.sleep(0.05)
    await task
    return seen


CONTACT = {"name": "李四", "phone": "13900000000", "email": "lisi@example.com"}


# ---------------------------------------------------------------------------
# 一、实用新型分支：脚本化产物
# ---------------------------------------------------------------------------

U_TITLE = "一种带卡扣式散热鳍片的电控盒壳体装置"
U_WORKING_TITLE = "卡扣式散热鳍片电控盒壳体"
U_TERMS = ["散热鳍片", "卡扣"]
U_PARTS = ["壳体", "散热鳍片", "卡扣"]
U_USER_UNCERTAIN = "卡扣弹性行程未标注"

U_MATERIAL_MD = f"""# 电控盒壳体结构改进说明

本文档由{SENSITIVE[0]}结构组编写，样机交付给{SENSITIVE[1]}试装，一期模具费 {SENSITIVE[2]}。

## 1. 现有结构与问题
现有电控盒壳体用四颗螺钉固定顶盖，拆装需专用工具，且顶盖为平板，散热面积不足。

## 2. 改进要点
- 壳体顶面与散热鳍片一体成型，鳍片沿长边阵列排布；
- 壳体两侧壁开口处对称设置弹性卡扣，顶盖免螺钉快拆；
- 卡扣与鳍片共面布置，整机高度不增加。
"""

U_MATERIAL_DIGEST = {
    "file": "material.md",
    "priority": 2,
    "summary_zh": "材料给出电控盒壳体的现有螺钉固定结构问题与鳍片一体成型、侧壁卡扣两处构造改进。",
    "tech_points": ["顶面与鳍片一体成型", "侧壁弹性卡扣免螺钉快拆", "卡扣与鳍片共面布置"],
    "components": ["壳体", "散热鳍片", "卡扣"],
    "figures_mentioned": ["总装线稿"],
    "sensitive_hits": SENSITIVE,
    "type_signals": {"invention": 0.1, "utility_model": 0.9, "design": 0.0},
}

U_POINTS_MD = """### 候选结构点分析

**U1 一体成型鳍片与侧壁卡扣的壳体**
- 技术背景：现有壳体螺钉固定、顶盖为平板。
- 构造创新：顶面与散热鳍片一体成型，侧壁开口处设弹性卡扣。
- 图证：总装线稿可见鳍片阵列与卡扣位置。
- 与现有技术区别：免螺钉快拆且散热面积增大。
- 可实施性：压铸模具即可实现。

**U2 卡扣与鳍片共面布置**
- 技术背景：附加卡扣通常凸出壳体轮廓。
- 构造创新：卡扣顶面与鳍片顶面齐平。
- 图证：总装线稿的侧向轮廓。
- 与现有技术区别：整机高度不增加。
- 可实施性：调整模具分型面即可。

```json
{
  "candidates": [
    {"id": "U1", "title": "一体成型鳍片与侧壁卡扣的壳体",
     "background": "现有壳体螺钉固定。", "innovation": "顶面与鳍片一体成型并设侧壁卡扣。",
     "distinction": "免螺钉快拆且散热面积增大。", "feasibility": "压铸模具即可实现。", "score": 86},
    {"id": "U2", "title": "卡扣与鳍片共面布置",
     "background": "卡扣通常凸出轮廓。", "innovation": "卡扣顶面与鳍片顶面齐平。",
     "distinction": "整机高度不增加。", "feasibility": "调整分型面即可。", "score": 71}
  ],
  "recommended": "U1",
  "fusion_note": "以 U1 为主线并入 U2 的共面布置。",
  "type_mismatch_note": ""
}
```
"""

U_STRUCTURE_SCHEMA = {
    "version": 1,
    "mode": "disclosure",
    "source_images": [LINEART],
    "parts": [
        {"id": "1", "name": "壳体", "shape": "矩形盒状", "material_hint": "铝合金压铸"},
        {"id": "2", "name": "散热鳍片", "shape": "沿长边排布的阵列鳍片", "material_hint": "unknown"},
        {"id": "3", "name": "卡扣", "shape": "弹性钩状", "material_hint": "unknown"},
    ],
    "relations": [
        {"from": "1", "to": "2", "type": "一体成型", "where": "壳体顶面", "seen_in": [LINEART]},
        {"from": "1", "to": "3", "type": "卡扣", "where": "壳体两侧壁开口处", "seen_in": [LINEART]},
    ],
    "spatial": ["散热鳍片位于壳体顶面并沿长边阵列排布", "卡扣对称设置于壳体两侧壁"],
    "function_of_structure": ["鳍片阵列扩大散热面积", "弹性卡扣实现免螺钉快拆"],
    "delta_hypothesis": ["卡扣顶面与鳍片顶面齐平，整机高度不增加"],
    "uncertain": ["鳍片间距与壁厚在材料中未标注"],
    "not_utility_model_signals": [],
}

U_FIGURE_PLAN = {
    "version": 1,
    "patent_type": "utility_model",
    "theme_summary": "电控盒壳体的鳍片与卡扣构造",
    "figures": [
        {
            "fig": 1, "role": "assembly", "path": LINEART, "covers": ["1", "2", "3"],
            "kind": "lineart", "relevance": 88, "quality": 90, "score": 89,
            "use_in_disclosure": True, "reason": "总装线稿，件号可与部件表对齐", "relates_to": [],
        },
        {
            "fig": None, "role": "perspective", "path": PHOTO_CLEAN, "covers": ["1"],
            "kind": "photo_clean", "relevance": 80, "quality": 82, "score": 81,
            "use_in_disclosure": True, "reason": "白底样机实拍", "relates_to": [],
        },
        {
            "fig": None, "role": "reference", "path": PHOTO_SCENE, "covers": [],
            "kind": "photo_scene", "relevance": 40, "quality": 45, "score": 42,
            "use_in_disclosure": False, "reason": "展台场景照，主体被遮挡", "relates_to": [],
        },
        {
            "fig": None, "role": "reference", "path": CAD, "covers": ["1", "2"],
            "kind": "cad", "relevance": 86, "quality": 88, "score": 87,
            "use_in_disclosure": False, "reason": "CAD 投影，仅作生图材料", "relates_to": [],
        },
    ],
}

U_SEARCH_TERMS = {
    "blocks": ["卡扣式散热鳍片", "电控盒壳体", "免螺钉快拆", "鳍片阵列"],
    "type_param": "utility_model",
    "rationale": "覆盖本案的鳍片构造、壳体对象与卡扣快拆三个方案要素。",
}

U_PREVIEW_MD = """### 交底书摘要预览

- **选定名称**：一种带卡扣式散热鳍片的电控盒壳体装置
- **专利类型**：实用新型
- **解决的技术问题**：螺钉固定拆装不便；平板顶盖散热面积不足。
- **关键部件与连接**：壳体、散热鳍片（一体成型）、卡扣（侧壁开口处）。
- **与最相近现有技术的区别**：鳍片与壳体一体成型，侧壁卡扣免螺钉快拆（待查新）。
"""

U_SKELETON = {
    "case_title": U_TITLE,
    "title_core_terms": U_TERMS,
    "modules": U_PARTS,
    "steps": [
        {"sid": "S1", "label": "对准壳体侧壁开口"},
        {"sid": "S2", "label": "压入卡扣完成锁合"},
        {"sid": "S3", "label": "反向按压卡扣拆卸顶盖"},
    ],
    "chapter_outline": [
        "一、现有结构与缺点", "二、所要解决的技术问题", "3.1 总体构成", "3.2 连接与配合",
        "3.3 空间布局与附图说明", "3.4 结构作用与使用或拆装过程", "3.5 关键尺寸或材料",
        "四、优点", "五、欲保护点", "六、其它",
    ],
    "terms_init": [
        {
            "term": "散热鳍片",
            "definition": "壳体顶面沿长边阵列排布、用于扩大散热面积的片状构造",
            "forbidden_variants": ["散热片条"],
            "source_section": "3.1",
        },
        {
            "term": "卡扣",
            "definition": "壳体侧壁开口处的弹性钩状构造，用于免螺钉锁合顶盖",
            "source_section": "3.2",
        },
    ],
}

TERMS_DELTA = """```json
{"terms_delta": {"add": [{"term": "免螺钉快拆", "definition": "无需螺钉与工具即可完成顶盖装拆", "source_section": "3.4"}], "update": []}}
```"""

U_CH1 = """## 一、相关技术背景与最接近现有技术的缺点

### 1.1 现有技术

检索说明：本次未进行系统性专利检索，以下现有结构依据材料中已知的常见做法概述，面向代理人陈述。

**方向一：螺钉固定的平板顶盖壳体**

- 结构方案：顶盖与壳体四角以螺钉固定，顶面为平板。
- 应用场景：工业电控盒的常规封装。
- 局限性：拆装需专用工具，且平板顶面散热面积不足。

检索总结：上述常见结构未把散热鳍片与壳体一体成型，也未在侧壁设置免螺钉的卡扣构造。

### 1.2 现有技术存在的缺点

- 螺钉固定导致拆装工序多，现场维护需专用工具，装配效率低；
- 平板顶盖散热面积不足，长时间运行时腔内温度偏高；
- 附加散热件与壳体分体装配，接触热阻大且占用高度。

## 二、本实用新型所要解决的技术问题

- 如何在不增加整机高度的前提下扩大壳体顶面的散热面积；
- 如何取消螺钉、实现顶盖的免工具装拆；
- 如何减少散热构造与壳体之间的接触热阻。

""" + TERMS_DELTA + "\n"

U_CH3 = """## 三、技术方案的详细阐述

### 3.1 总体构成

| 件号 | 名称 | 形状要点 |
|------|------|----------|
| 1 | 壳体 | 矩形盒状，顶面平整 |
| 2 | 散热鳍片 | 沿长边排布的阵列鳍片 |
| 3 | 卡扣 | 弹性钩状 |

本装置由壳体、散热鳍片与卡扣三个部件构成：壳体承载电控器件，散热鳍片位于壳体顶面，
卡扣设置在壳体两侧壁，三者共同构成免螺钉装拆的电控盒结构。

### 3.2 连接与配合

| 自 | 至 | 类型 | 位置 |
|------|------|------|------|
| 壳体 | 散热鳍片 | 一体成型 | 壳体顶面 |
| 壳体 | 卡扣 | 卡扣 | 壳体两侧壁开口处 |

壳体与散热鳍片以一体成型方式连接，鳍片根部与壳体顶面之间没有装配界面，热量可直接由壳体
传导至鳍片；壳体与卡扣的配合为卡扣式咬合，卡扣的弹性钩部伸入侧壁开口并在压入到位后回弹限位，
反向按压即可解除锁合。

### 3.3 空间布局与附图说明

散热鳍片位于壳体顶面并沿长边阵列排布，卡扣对称设置于壳体两侧壁，鳍片顶面与卡扣顶面齐平。

如图 1，示出本装置的总装形态：可见壳体顶面的散热鳍片阵列，以及两侧壁开口处的卡扣位置。
鳍片间距与壁厚尚待确认，图中仅示意其排布关系。

""" + TERMS_DELTA + "\n"

U_CH35 = """### 3.4 结构作用与使用或拆装过程

- 散热鳍片与壳体一体成型，消除了分体装配的接触热阻，解决腔内温度偏高的问题；
- 卡扣位于壳体侧壁开口处，以弹性钩部限位替代螺钉，解决拆装需专用工具的问题；
- 卡扣顶面与散热鳍片顶面齐平，解决附加构造占用整机高度的问题。

S1，对准壳体侧壁开口：把顶盖沿壳体长边方向对准两侧壁的开口位置。

S2，压入卡扣完成锁合：向下压入顶盖，卡扣的弹性钩部沿开口斜面变形并在到位后回弹限位。

S3，反向按压卡扣拆卸顶盖：向壳体内侧按压卡扣使钩部脱离开口，即可取下顶盖。

### 3.5 关键尺寸或材料

| 部位 | 材料或尺寸 | 依据 |
|------|------------|------|
| 壳体 | 铝合金压铸 | 事实合同 parts.material_hint |
| 散热鳍片 | 待确认 | 材料未写明 |
| 卡扣 | 待确认 | 材料未写明 |

""" + TERMS_DELTA + "\n"

# 首版第五章故意写成方法书式 → section5_device_lint + AUDIT 双查打回 → u4.repair1
U_CH456_BAD = """## 四、与现有技术相比的优点

本实用新型以一体成型的散热鳍片与侧壁卡扣替代分体散热件与螺钉固定，兼顾散热与装拆效率。

- 散热面积增大：鳍片沿壳体长边阵列排布，散热面积较平板顶盖显著增加；
- 装拆效率提升：卡扣免螺钉锁合，现场维护无需专用工具即可开合顶盖；
- 高度不增加：卡扣顶面与鳍片顶面齐平，附加构造不占用整机高度。

## 五、本实用新型的技术关键点与欲保护点

- 一种电控盒散热方法，包括以下步骤：先对准侧壁开口，再压入顶盖完成锁合，最后按压卡扣拆卸；
- 步骤一：把散热鳍片按阵列方式布置在壳体顶面，使散热面积大于平板顶盖的散热面积；
- 步骤二：把卡扣压入壳体侧壁开口并回弹限位，使顶盖在不使用螺钉的情况下保持锁合。

## 六、其它

**实施例一：标准规格电控盒的装配**：按 S1 至 S2 完成顶盖锁合，散热鳍片沿壳体长边阵列排布，
卡扣钩部在开口内回弹限位，装配过程无需螺钉与专用工具。

**实施例二：现场维护时的顶盖拆卸**：按 S3 向内按压两侧卡扣使钩部脱离开口，顶盖即可整体取下，
维护完成后重复 S1 与 S2 复位；上述规格仅为示例，不作为权利要求限制。

- 技术效果：一体成型消除接触热阻，腔内温升下降，且顶盖装拆无需工具与螺钉；
- 技术效果：卡扣与散热鳍片共面布置，整机高度与原平板顶盖方案保持一致。

""" + TERMS_DELTA + "\n"

U_CH456_FIXED = U_CH456_BAD.replace(
    """- 一种电控盒散热方法，包括以下步骤：先对准侧壁开口，再压入顶盖完成锁合，最后按压卡扣拆卸；
- 步骤一：把散热鳍片按阵列方式布置在壳体顶面，使散热面积大于平板顶盖的散热面积；
- 步骤二：把卡扣压入壳体侧壁开口并回弹限位，使顶盖在不使用螺钉的情况下保持锁合。""",
    """- 一种电控盒壳体装置，包括壳体、散热鳍片与卡扣；其特征在于，散热鳍片与壳体顶面一体成型并沿
  长边阵列排布，卡扣设置于壳体两侧壁开口处并以弹性钩部与开口咬合限位；
- 从属构造：卡扣的钩部顶面与散热鳍片的顶面齐平，使附加构造不超出壳体顶面所在的轮廓；
- 从属构造：散热鳍片的根部与壳体顶面之间无装配界面，热量由壳体直接传导至鳍片根部。""",
)

U_CLAIM_AUDIT_FAIL = {
    "verdict": "fail",
    "is_device_form": False,
    "problems": ["主保护点写成了「一种……方法，包括以下步骤」，不是装置书式"],
    "suggestion": "改写为「一种电控盒壳体装置，包括……；其特征在于……」",
}
U_CLAIM_AUDIT_PASS = {"verdict": "pass", "is_device_form": True, "problems": [], "suggestion": ""}

SELF_CHECK_PASS = {"verdict": "pass", "patches": [], "unresolved": []}

U_CLAIM_BIAS = {
    "groups": [
        {
            "axis": "构造限定侧重 vs 配合关系侧重",
            "option_a": {
                "label": "更偏构造限定：以鳍片阵列的形状特征为保护重心",
                "basis_quote": "散热鳍片与壳体顶面一体成型",
            },
            "option_b": {
                "label": "更偏配合关系：以卡扣与开口的咬合限位为保护重心",
                "basis_quote": "卡扣设置于壳体两侧壁开口处",
            },
        }
    ]
}


def _utility_llm() -> FakeLLM:
    return FakeLLM(
        stream={
            "recap": "- 技术主题：电控盒壳体结构改进\n- 专利类型：实用新型\n- 技术联系人：李四",
            "points": U_POINTS_MD,
            "preview": U_PREVIEW_MD,
            "u1": U_CH1,
            "u2": U_CH3,
            "u3": U_CH35,
            "u4": U_CH456_BAD,
        },
        chat={"u4.repair1": U_CH456_FIXED},
        struct_={
            "digest.0": U_MATERIAL_DIGEST,
            "schema": U_STRUCTURE_SCHEMA,
            "figure_plan": U_FIGURE_PLAN,
            "terms": U_SEARCH_TERMS,
            "g0": U_SKELETON,
            "u4.claim_audit1": U_CLAIM_AUDIT_FAIL,
            "u4.claim_audit2": U_CLAIM_AUDIT_PASS,
            "audit": SELF_CHECK_PASS,
            "claim_bias": U_CLAIM_BIAS,
        },
    )


# ---------------------------------------------------------------------------
# 二、外观设计分支：脚本化产物
# ---------------------------------------------------------------------------

D_PRODUCT = "折叠台灯"
D_POINTS = ["灯臂与底座的折线形转轴造型", "底座圆环凹槽装饰"]

D_MATERIAL_MD = f"""# 折叠台灯外观设计说明

本设计由{SENSITIVE[0]}工业设计组完成，首批样机交{SENSITIVE[1]}试销，开模预算 {SENSITIVE[2]}。

## 1. 造型描述
灯臂由两段直杆经转轴相连，展开后呈折线形；底座为圆盘，正面设一圈等距凹槽。

## 2. 配色
主体哑光白，底座外沿浅木色。
"""

D_MATERIAL_DIGEST = {
    "file": "material.md",
    "priority": 2,
    "summary_zh": "材料给出折叠台灯的整体造型、转轴折线形态、底座圆环凹槽装饰与配色。",
    "tech_points": ["折线形灯臂", "圆盘底座", "等距凹槽装饰"],
    "components": ["灯臂", "底座"],
    "figures_mentioned": ["立体线稿", "白底实拍"],
    "sensitive_hits": SENSITIVE,
    "type_signals": {"invention": 0.0, "utility_model": 0.1, "design": 0.9},
}

D_POINTS_MD = """### 候选外观点分析

**D1 折线形灯臂与圆环凹槽底座**
- 产品名称与用途：折叠台灯，用于桌面照明。
- 设计要点：灯臂与底座的折线形转轴造型；底座圆环凹槽装饰。
- 视图与图证：立体线稿与白底实拍可见灯臂折线与底座凹槽。
- 与在先外观差异：常见台灯多为直杆造型。

```json
{
  "candidates": [
    {"id": "D1", "title": "折叠台灯",
     "background": "常见台灯多为直杆造型。", "innovation": "灯臂呈折线并配圆环凹槽底座。",
     "distinction": "轮廓与装饰线条明显不同。", "feasibility": "注塑与车削可实现。", "score": 84}
  ],
  "recommended": "D1",
  "fusion_note": "以 D1 单点推进。",
  "type_mismatch_note": ""
}
```
"""

D_APPEARANCE_SCHEMA = {
    "version": 1,
    "mode": "disclosure",
    "source_images": [LINEART, PHOTO_CLEAN],
    "product_name": D_PRODUCT,
    "overall_shape": "折线形灯臂与圆盘底座构成的可折叠桌面灯具",
    "product_form": "solid",
    "claimed_faces": ["主视", "俯视"],
    "omitted_views": [{"name": "仰视", "reason": "底面为平整胶垫，无设计要点"}],
    "views": [
        {"name": "立体图", "notes": "展开状态的整体折线轮廓", "source_image": LINEART},
        {"name": "主视图", "notes": "灯臂与底座的高度比例", "source_image": PHOTO_CLEAN},
    ],
    "ornament": ["底座正面的一圈等距凹槽"],
    "color": ["主体哑光白", "底座外沿浅木色"],
    "design_points": D_POINTS,
    "contrast_to_prior": ["常见台灯多为直杆造型，本设计灯臂呈折线（假设）"],
    "uncertain": ["底座外沿的具体色号在材料中看不清"],
    "not_design_signals": [],
}

D_FIGURE_PLAN = {
    "version": 1,
    "patent_type": "design",
    "theme_summary": "折叠台灯的折线灯臂与圆环凹槽底座",
    "figures": [
        {
            "fig": 1, "role": "perspective", "path": LINEART, "covers": ["立体图"],
            "kind": "lineart", "relevance": 88, "quality": 90, "score": 89,
            "use_in_disclosure": True, "reason": "立体线稿，轮廓清晰", "relates_to": [],
        },
        {
            "fig": 2, "role": "ortho", "path": PHOTO_CLEAN, "covers": ["主视"],
            "kind": "photo_clean", "relevance": 84, "quality": 86, "score": 85,
            "use_in_disclosure": True, "reason": "白底干净实拍，可见灯臂比例", "relates_to": [],
        },
        {
            "fig": None, "role": "reference", "path": PHOTO_SCENE, "covers": [],
            "kind": "photo_scene", "relevance": 40, "quality": 45, "score": 42,
            "use_in_disclosure": False, "reason": "场景宣传照，含无关陈设", "relates_to": [],
        },
        {
            "fig": None, "role": "reference", "path": CAD, "covers": [],
            "kind": "cad", "relevance": 86, "quality": 88, "score": 87,
            "use_in_disclosure": False, "reason": "CAD 投影，仅作生图材料", "relates_to": [],
        },
    ],
}

D_SEARCH_TERMS = {
    "blocks": ["折叠台灯", "折线形灯臂", "圆环凹槽底座", "桌面灯具外观"],
    "type_param": "design",
    "rationale": "覆盖产品对象、灯臂造型与底座装饰三个要素。",
}

D_PREVIEW_MD = """### 交底底稿摘要预览

- **选定名称**：折叠台灯
- **专利类型**：外观设计
- **设计要点**：灯臂与底座的折线形转轴造型；底座圆环凹槽装饰。
- **视图清单**：立体图、主视图、俯视图（仰视省略）。
- **与在先外观的区别**：常见台灯多为直杆造型（待查新）。
"""

D_SKELETON = {
    "case_title": D_PRODUCT,
    "title_core_terms": ["台灯"],
    "modules": D_POINTS,
    "steps": [
        {"sid": "S1", "label": "立体图"},
        {"sid": "S2", "label": "主视图"},
        {"sid": "S3", "label": "俯视图"},
    ],
    "chapter_outline": ["一、产品名称与用途", "二、设计要点", "三、视图说明", "四、与在先外观的差异", "五、其它"],
    "terms_init": [
        {"term": "灯臂", "definition": "由两段直杆经转轴相连、展开后呈折线的支撑构件", "source_section": "二"},
        {"term": "底座", "definition": "承托灯臂的圆盘状部件，正面设一圈等距凹槽", "source_section": "二"},
    ],
}

D_CH123 = """## 一、产品名称与用途

本设计的产品名称为折叠台灯，用于桌面照明；使用时展开灯臂调节照射角度，收纳时折叠贴合底座。

## 二、设计要点

- 灯臂与底座的折线形转轴造型：灯臂由两段直杆经转轴相连，展开后整体呈折线轮廓；
- 底座圆环凹槽装饰：底座正面沿外沿设一圈等距凹槽，形成连续的环形装饰线条；
- 配色为主体哑光白、底座外沿浅木色，底座外沿的具体色号待确认。

## 三、视图说明

- 立体图：见图 1，示出折叠台灯展开状态的整体折线轮廓与底座圆环凹槽的相对位置；
- 主视图：见图 2，示出灯臂两段直杆与底座的高度比例，以及转轴处的折线转折关系；
- 俯视：自上方观察可见底座圆环凹槽的等距排布与灯臂在底座上的投影位置；
- 简要说明：仰视省略，底面为平整胶垫，无设计要点，故不单独出图。

""" + TERMS_DELTA + "\n"

D_CH45 = """## 四、与在先外观的主要差异

- 轮廓差异：见图 1，本设计的灯臂展开后呈折线轮廓，而常见台灯多为直杆造型（该对比为依据材料的假设）；
- 装饰差异：见图 2 的底座部位，本设计沿底座外沿设一圈等距凹槽，形成连续环形线条，常见底座多为素面；
- 比例差异：灯臂两段直杆的长度比例与底座直径的比值明显不同于常见的等长直杆造型。

本次未进行系统性专利检索，上述对比仅依据材料中已知的常见造型，未引用任何在先公开文件。

## 五、其它

- 色彩说明：主体哑光白、底座外沿浅木色；底座外沿的具体色号待确认，不作为设计要点写死；
- 使用状态参考图：折叠收纳状态可另行提供参考图，本稿不作为要点落面单独出图。

""" + TERMS_DELTA + "\n"

D_CLAIM_BIAS = {
    "groups": [
        {
            "axis": "整体轮廓侧重 vs 装饰线条侧重",
            "option_a": {
                "label": "更偏整体轮廓：以灯臂折线与底座的比例关系为重心",
                "basis_quote": "灯臂由两段直杆经转轴相连，展开后整体呈折线轮廓",
            },
            "option_b": {
                "label": "更偏装饰线条：以底座环形凹槽的连续线条为重心",
                "basis_quote": "底座正面沿外沿设一圈等距凹槽",
            },
        }
    ]
}


def _design_llm() -> FakeLLM:
    return FakeLLM(
        stream={
            "recap": "- 技术主题：折叠台灯外观\n- 专利类型：外观设计\n- 技术联系人：李四",
            "points": D_POINTS_MD,
            "preview": D_PREVIEW_MD,
            "d1": D_CH123,
            "d2": D_CH45,
        },
        struct_={
            "digest.0": D_MATERIAL_DIGEST,
            "schema": D_APPEARANCE_SCHEMA,
            "figure_plan": D_FIGURE_PLAN,
            "terms": D_SEARCH_TERMS,
            "g0": D_SKELETON,
            "audit": SELF_CHECK_PASS,
            "claim_bias": D_CLAIM_BIAS,
        },
    )


# ---------------------------------------------------------------------------
# 门控答复
# ---------------------------------------------------------------------------

_GATE_DEFAULTS: dict[str, Any] = {}


def _schema_fill_answer(pending: dict[str, Any]) -> dict[str, Any]:
    """A3b 回填：改一处事实合同 + 故意把 CAD 勾成入文（服务端硬规则必须打回）。"""
    if pending.get("kind") != "schema_fill":
        return {}
    default = pending["default"]
    _GATE_DEFAULTS[default["patent_type"]] = default
    schema = json.loads(json.dumps(default["schema"]))
    schema["uncertain"] = [*(schema.get("uncertain") or []), U_USER_UNCERTAIN]
    plan = json.loads(json.dumps(default["figure_plan"]))
    for item in plan["figures"]:
        if item.get("kind") == "cad":
            item["use_in_disclosure"] = True
    return {"schema": schema, "figure_plan": plan, "note": "已补充一条待确认项"}


def _answers(patent_type: str, topic: str, working_title: str, point_ids: list[str]) -> dict[str, Any]:
    return {
        "intake": {
            "topic": topic,
            "patent_type": patent_type,
            "contact": dict(CONTACT),
            "notes": "",
        },
        "points_mining": {
            "selected_ids": point_ids,
            "working_title": working_title,
            "instruction": "",
        },
        "schema_fill": _schema_fill_answer,
        "prior_art_search": {"action": "skip", "reason": "本次先出稿，检索另行安排"},
        "preview": {"action": "confirm", "feedback": ""},
        "deliver": {"skip": True, "choices": [], "apply": False, "note": ""},
    }


# ---------------------------------------------------------------------------
# 1. 实用新型分支全流水线
# ---------------------------------------------------------------------------

_um: dict[str, Any] = {}
_ds: dict[str, Any] = {}


async def test_utility_pipeline_runs_to_deliver(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """实用新型九步跑到 deliver：schema_fill 门控在挖点之后、查新之前。"""
    from app.pipelines import engine

    fake = _install_llm(monkeypatch, _utility_llm())
    _install_failing_search(monkeypatch)

    case_id = _new_case(client, "交底书实用新型测试", "utility_model")
    file_id = _upload(client, case_id, "material.md", U_MATERIAL_MD.encode("utf-8"), "text/markdown")
    _upload_images(client, case_id)

    steps = _steps(case_id)
    assert [s.key for s in steps][3] == "schema_fill" and len(steps) == 9

    task = engine.start(
        case_id,
        steps,
        run_group="initial",
        start_payload={"file_ids": [file_id], "patent_type": "utility_model"},
    )
    seen = await asyncio.wait_for(
        _drive(case_id, task, _answers("utility_model", "电控盒壳体结构改进", U_WORKING_TITLE, ["U1", "U2"])),
        timeout=1200,
    )

    assert _case_row(case_id)["status"] == "completed", _failures(case_id)
    _um["case_id"] = case_id
    _um["state"] = _case_state(case_id)
    _um["calls"] = list(fake.calls)
    _um["seen"] = seen
    _um["default"] = _GATE_DEFAULTS.get("utility_model")

    keys = [s["step_key"] for s in seen]
    assert "schema_fill" in keys, f"未经过 A3b 门控：{keys}"
    assert keys.index("schema_fill") < keys.index("prior_art_search")
    assert {"step_key": "schema_fill", "kind": "schema_fill"} in seen

    # 第五章书式 lint 生效：首版方法书式 → 一次章节 REPAIR
    assert "chat:u4.repair1" in fake.calls, "第五章方法书式未触发装置书式 REPAIR"
    assert "struct:u4.claim_audit1" in fake.calls and "struct:u4.claim_audit2" in fake.calls


def test_utility_schema_fill_state() -> None:
    """A3b 产物落 state：用户回填生效；CAD / 场景实拍 / 实拍均被服务端规则挡在文外。"""
    from app.services import figure_plan as figure_plan_service

    state = _um["state"]
    schema = state["structure_schema"]
    assert [p["name"] for p in schema["parts"]] == U_PARTS
    assert schema["relations"][0]["from"] == "1", "relations 应按合同字段名 `from` 落库"
    assert U_USER_UNCERTAIN in schema["uncertain"], "用户在卡片里补的待确认项未落进 state"

    card = state["schema_fill"]
    assert card["patent_type"] == "utility_model" and card["skipped"] is False
    assert card["edited"] is True and card["note"] == "已补充一条待确认项"
    assert card["type_suggest"] is None

    plan = state["figure_plan"]
    by_path = {f["path"]: f for f in plan["figures"]}
    assert len(by_path) == 4
    used = figure_plan_service.used_figures(plan)
    assert [f["path"] for f in used] == [LINEART], "实用新型入文只收合格线稿"
    assert used[0]["fig"] == 1 and used[0]["kind"] == "lineart"
    assert by_path[CAD]["use_in_disclosure"] is False, "CAD 永不入文（用户勾选也须打回）"
    assert "CAD" in by_path[CAD]["reason"]
    assert by_path[PHOTO_SCENE]["use_in_disclosure"] is False
    assert "低于合格线" in by_path[PHOTO_SCENE]["reason"]
    assert by_path[PHOTO_CLEAN]["use_in_disclosure"] is False, "实用新型不收实拍"

    # 用户误勾的 CAD 被规则改回，须留下调整记录
    adjusted = {a["path"] for a in card["figure_report"]["adjusted"]}
    assert CAD in adjusted

    # 有合格线稿 → 不需要下发线稿绘制说明
    assert card["lineart_brief"]["needed"] is False


def test_utility_document_structure() -> None:
    """实用新型终稿：文头 + 六章 + 3.1/3.2/3.3 小节；第五章为装置书式。"""
    from app.services import disclosure_build as build_service

    state = _um["state"]
    text = state["final_markdown"]

    assert text.lstrip().startswith("# 技术交底书")
    assert f"**案件名称**：{U_TITLE}" in text
    assert "**专利类型**：实用新型" in text
    assert "## 注意事项" in text
    assert f"- 姓名：{CONTACT['name']}" in text

    for heading in ("## 一、", "## 二、", "## 三、", "## 四、", "## 五、", "## 六、"):
        assert heading in text, f"终稿缺少章节：{heading}"
    for section in ("### 3.1 总体构成", "### 3.2 连接与配合", "### 3.3 空间布局与附图说明",
                    "### 3.4 结构作用与使用或拆装过程", "### 3.5 关键尺寸或材料"):
        assert section in text, f"终稿缺少小节：{section}"

    # 部件与连接类型逐条落在正文
    for name in U_PARTS:
        assert name in text
    assert "一体成型" in text and "其特征在于" in text

    # 第五章装置书式 lint 通过，且不再有方法书式主保护点
    assert build_service.section5_device_lint(text) == []
    assert "一种电控盒散热方法" not in text

    # 入文附图确定性嵌入（图题 + 图片引用）
    assert "#### 附图" in text and "图1 总装图（线稿）" in text
    assert LINEART in text and CAD not in text and PHOTO_SCENE not in text

    # 未检索时如实写明；无编造链接；无元信息与敏感词
    assert state["prior_art"]["skipped"] is True and state["prior_art"]["searched"] is False
    assert build_service.url_lint(text, set()) == []
    assert build_service.meta_leak_lint(text) == []
    assert "terms_delta" not in text
    for term in SENSITIVE:
        assert term not in text, f"终稿泄漏敏感词：{term}"

    # 章节报告：仅第四五六章因书式打回一次
    by_key = {r["key"]: r for r in state["build_report"]["chapters"]}
    assert by_key["u4"]["repairs"] == 1
    assert all(r["repairs"] == 0 for k, r in by_key.items() if k != "u4")
    assert state["build_report"]["unresolved"] == []
    assert state["build_report"]["figures"] == 1


def test_utility_docx_artifact() -> None:
    """实用新型交付物：md + docx 各一版，Word 可打开且含各章标题。"""
    from docx import Document

    case_id = _um["case_id"]
    mds = _artifacts(case_id, "disclosure_md")
    assert len(mds) == 1 and mds[0]["version"] == 1
    assert Path(mds[0]["stored_path"]).read_text(encoding="utf-8") == _um["state"]["final_markdown"]
    assert mds[0]["summary"] == "交底书定稿（实用新型）"

    docxs = _artifacts(case_id, "disclosure_docx")
    assert len(docxs) == 1, f"docx 未产出：{_um['state']['deliver']['files'].get('docx_error')}"
    blob = "\n".join(p.text.strip() for p in Document(docxs[0]["stored_path"]).paragraphs)
    assert "技术交底书" in blob
    for heading in ("一、相关技术背景", "三、技术方案的详细阐述", "3.1 总体构成",
                    "3.2 连接与配合", "五、本实用新型的技术关键点与欲保护点"):
        assert heading in blob, f"DOCX 缺少章节：{heading}"


# ---------------------------------------------------------------------------
# 2. 外观设计分支全流水线
# ---------------------------------------------------------------------------


async def test_design_pipeline_runs_to_deliver(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """外观九步跑到 deliver：实拍与线稿都入文，视图说明按要点落面写。"""
    from app.pipelines import engine

    fake = _install_llm(monkeypatch, _design_llm())
    _install_failing_search(monkeypatch)

    case_id = _new_case(client, "交底书外观设计测试", "design")
    file_id = _upload(client, case_id, "material.md", D_MATERIAL_MD.encode("utf-8"), "text/markdown")
    _upload_images(client, case_id)

    task = engine.start(
        case_id,
        _steps(case_id),
        run_group="initial",
        start_payload={"file_ids": [file_id], "patent_type": "design"},
    )
    seen = await asyncio.wait_for(
        _drive(case_id, task, _answers("design", "折叠台灯外观造型", D_PRODUCT, ["D1"])),
        timeout=1200,
    )

    assert _case_row(case_id)["status"] == "completed", _failures(case_id)
    _ds["case_id"] = case_id
    _ds["state"] = _case_state(case_id)
    _ds["calls"] = list(fake.calls)
    _ds["default"] = _GATE_DEFAULTS.get("design")

    keys = [s["step_key"] for s in seen]
    assert keys.index("schema_fill") < keys.index("prior_art_search")
    assert {"step_key": "schema_fill", "kind": "schema_fill"} in seen


def test_design_schema_fill_state() -> None:
    """外观 A3b：事实合同落 state；实拍 + 线稿都入文，场景照与 CAD 不入文。"""
    from app.services import figure_plan as figure_plan_service

    state = _ds["state"]
    schema = state["appearance_schema"]
    assert schema["product_name"] == D_PRODUCT
    assert schema["claimed_faces"] == ["主视", "俯视"]
    assert schema["design_points"] == D_POINTS
    assert U_USER_UNCERTAIN in schema["uncertain"], "用户回填未覆盖 state"

    plan = state["figure_plan"]
    by_path = {f["path"]: f for f in plan["figures"]}
    used = figure_plan_service.used_figures(plan)
    assert [f["path"] for f in used] == [LINEART, PHOTO_CLEAN], "外观入文＝线稿 + 干净实拍"
    assert [f["fig"] for f in used] == [1, 2]
    assert by_path[CAD]["use_in_disclosure"] is False, "CAD 永不入文（用户勾选也须打回）"
    assert by_path[PHOTO_SCENE]["use_in_disclosure"] is False

    card = state["schema_fill"]
    assert card["patent_type"] == "design" and card["used_figures"] == 2
    assert card["lineart_brief"]["needed"] is False


def test_design_document_structure() -> None:
    """外观终稿：产品名称文头 + 五章，三、视图说明按 claimed_faces 展开。"""
    from app.services import disclosure_build as build_service

    state = _ds["state"]
    text = state["final_markdown"]

    assert text.lstrip().startswith("# 外观设计说明")
    assert f"**产品名称**：{D_PRODUCT}" in text
    assert "**专利类型**：外观设计" in text

    for heading in ("## 一、产品名称与用途", "## 二、设计要点", "## 三、视图说明",
                    "## 四、与在先外观的主要差异", "## 五、其它"):
        assert heading in text, f"终稿缺少章节：{heading}"
    for face in ("主视", "俯视", "仰视"):
        assert face in text, f"视图说明缺少「{face}」"
    assert "六视图" not in text

    # 实拍与线稿都进了附图块
    assert "#### 附图（实拍与线稿）" in text
    assert "图1 立体图（线稿）" in text and "图2 正投影图（干净实拍）" in text
    assert LINEART in text and PHOTO_CLEAN in text
    assert CAD not in text and PHOTO_SCENE not in text

    assert build_service.visible_only_lint(text) == []
    assert build_service.url_lint(text, set()) == []
    assert build_service.meta_leak_lint(text) == []
    assert "terms_delta" not in text
    for term in SENSITIVE:
        assert term not in text

    assert state["build_report"]["figures"] == 2
    assert state["build_report"]["unresolved"] == []


def test_design_docx_artifact() -> None:
    """外观交付物：md + docx，Word 内含「三、视图说明」。"""
    from docx import Document

    case_id = _ds["case_id"]
    mds = _artifacts(case_id, "disclosure_md")
    assert len(mds) == 1 and mds[0]["summary"] == "交底书定稿（外观设计）"

    docxs = _artifacts(case_id, "disclosure_docx")
    assert len(docxs) == 1, f"docx 未产出：{_ds['state']['deliver']['files'].get('docx_error')}"
    blob = "\n".join(p.text.strip() for p in Document(docxs[0]["stored_path"]).paragraphs)
    assert "外观设计说明" in blob
    for heading in ("一、产品名称与用途", "二、设计要点", "三、视图说明", "四、与在先外观的主要差异"):
        assert heading in blob, f"DOCX 缺少章节：{heading}"


# ---------------------------------------------------------------------------
# 3. 卡片契约与发明分支回归
# ---------------------------------------------------------------------------


def test_schema_fill_card_payload() -> None:
    """A3b 卡片：事实合同表单 + 逐图入文选择 + 图片清单，都随门控下发给前端。"""
    from app.pipelines import disclosure as pipeline

    default = _um["default"]
    assert default is not None, "未捕获到 schema_fill 门控的 default 载荷"
    assert default["patent_type"] == "utility_model"
    assert {"schema", "figure_plan", "lineart_brief", "warnings", "images"} <= set(default)
    assert [i["path"] for i in default["images"]] == [LINEART, PHOTO_CLEAN, PHOTO_SCENE, CAD]
    # 平台提示语里点名被规则挡下的图
    assert any(PHOTO_CLEAN in w for w in default["warnings"])

    form = pipeline.schema_fill_form("utility_model")
    assert "parts" in form["properties"]["schema"]["properties"]
    assert "figures" in form["properties"]["figure_plan"]["properties"]
    design_form = pipeline.schema_fill_form("design")
    assert "claimed_faces" in design_form["properties"]["schema"]["properties"]


def test_lineart_brief_when_no_lineart() -> None:
    """缺合格线稿时下发「线稿绘制说明」；外观仅有实拍时提示补图风险但不阻断。"""
    from app.services import figure_plan as figure_plan_service

    photo_only = {
        "version": 1,
        "patent_type": "design",
        "figures": [
            {
                "fig": 1, "role": "perspective", "path": PHOTO_CLEAN, "kind": "photo_clean",
                "relevance": 84, "quality": 86, "score": 85, "use_in_disclosure": True,
                "reason": "白底实拍", "relates_to": [], "covers": [],
            }
        ],
    }
    brief = figure_plan_service.lineart_brief(
        "design", plan=photo_only, product_name=D_PRODUCT, goals={"立体图": "整体折线轮廓"}
    )
    assert brief["needed"] is True
    assert brief["views"] and "黑白线稿" in brief["views"][0]["gen_prompt"]
    assert "存在补图风险" in brief["risk_note"]

    empty = {"version": 1, "patent_type": "utility_model", "figures": []}
    ubrief = figure_plan_service.lineart_brief("utility_model", plan=empty, product_name="电控盒壳体")
    assert ubrief["needed"] is True and ubrief["views"][0]["view_name"] == "总装立体图"
    assert "线稿" in ubrief["risk_note"]


async def test_invention_branch_unaffected(client: TestClient) -> None:
    """发明分支回归：仍是八步、无 schema_fill；handler 对发明直通不产事实合同。"""
    from types import SimpleNamespace

    from app.pipelines import disclosure as pipeline
    from app.pipelines import registry

    steps = registry.build_steps({"module": "disclosure", "patent_type": "invention"})
    assert len(steps) == 8 and "schema_fill" not in [s.key for s in steps]
    assert "invention" not in pipeline.SCHEMA_SPEC

    ctx = SimpleNamespace(case={"patent_type": "invention"}, state={}, case_id="not-used")
    assert await pipeline.schema_fill(ctx) == {}
